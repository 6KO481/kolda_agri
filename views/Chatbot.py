"""
PAGE 6 — CHATBOT IA
Agent conversationnel — Kolda Agri
Design : inspiré Claude.ai · Police Inter · Thème sombre
"""

import sys
import re
import io
import json
import time
import os
import tempfile
from pathlib import Path
from datetime import datetime

import streamlit as st
import pandas as pd

ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT / "db"))

from utils import get_connection, DB_PATH

try:
    import plotly.express as px
    PLOTLY_OK = True
except ImportError:
    PLOTLY_OK = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    MPL_OK = True
except ImportError:
    MPL_OK = False


# ══════════════════════════════════════════════════════════════
# CONFIG (couleurs & UI — sans tokens sensibles)
# ══════════════════════════════════════════════════════════════

CFG = {
    "body_bg":       "#0d1117",
    "surface_bg":    "#161b22",
    "surface_alt":   "#1c2128",
    "border":        "#30363d",
    "text_primary":  "#e6edf3",
    "text_secondary":"#8b949e",
    "text_muted":    "#6e7681",
    "accent":        "#3fb950",
    "accent_blue":   "#58a6ff",
    "accent_orange": "#d29922",
    "font":          "Inter",
    # Header identique à Carte.py
    "hdr_bg":        "#1c2a1e",
    "hdr_border":    "#3fb950",
    "hdr_text":      "#e6edf3",
    # Onglets identiques à Carte.py
    "tab_active":    "#3fb950",
    "subtab_active": "#58a6ff",
}

# ── Accès sécurisé aux secrets (Streamlit Cloud ou .streamlit/secrets.toml) ──
def _secret(key: str, fallback: str = "") -> str:
    """Lit un secret depuis st.secrets (Streamlit Cloud / secrets.toml).
    En développement local, créez .streamlit/secrets.toml avec :
        hf_token    = "hf_..."
        groq_token  = "gsk_..."
    """
    try:
        return st.secrets[key]
    except Exception:
        return fallback

SESSIONS_DIR = ROOT / "chat_sessions"
SESSIONS_DIR.mkdir(parents=True, exist_ok=True)


# ══════════════════════════════════════════════════════════════
# THÈME — identique Carte.py, config hardcodée
# ══════════════════════════════════════════════════════════════

def apply_theme():
    c = CFG

    def hex_rgba(h, a=0.07):
        h = h.lstrip("#")
        if len(h) == 6:
            r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
            return f"rgba({r},{g},{b},{a})"
        return f"rgba(63,185,80,{a})"

    tab_bg    = hex_rgba(c["tab_active"])
    subtab_bg = hex_rgba(c["subtab_active"])

    st.markdown(f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&display=swap');

    html, body, [class*="css"], p, span, div, button, input, textarea, .stMarkdown {{
        font-family: '{c["font"]}', -apple-system, BlinkMacSystemFont, sans-serif !important;
    }}
    .stApp {{ background-color: {c["body_bg"]} !important; }}
    .main .block-container {{
        padding-top: 0 !important; margin-top: 0 !important;
        max-width: 820px !important;
        padding-left: 1.2rem !important; padding-right: 1.2rem !important;
    }}
    header[data-testid="stHeader"] {{ height: 0 !important; overflow: hidden !important; }}

    /* Onglets principaux — identique Carte.py */
    [data-baseweb="tab-list"] {{
        gap: 0 !important; width: 100% !important;
        border-bottom: 1px solid {c["border"]} !important;
        background: transparent !important;
    }}
    [data-baseweb="tab"] {{
        flex: 1 1 0 !important; justify-content: center !important;
        padding: 10px 4px !important; font-size: 0.85rem !important;
        font-weight: 500 !important; color: {c["text_secondary"]} !important;
        border-bottom: 2px solid transparent !important;
        background: transparent !important;
    }}
    [data-baseweb="tab"][aria-selected="true"] {{
        border-bottom-color: {c["tab_active"]} !important;
        color: {c["tab_active"]} !important; background: {tab_bg} !important;
    }}
    [data-baseweb="tab"]:hover:not([aria-selected="true"]) {{
        background: rgba(255,255,255,0.03) !important;
        color: {c["text_primary"]} !important;
    }}
    .stTabs .stTabs [data-baseweb="tab"][aria-selected="true"] {{
        border-bottom-color: {c["subtab_active"]} !important;
        color: {c["subtab_active"]} !important; background: {subtab_bg} !important;
    }}

    /* Chat input — style Claude.ai */
    [data-testid="stChatInputContainer"] {{
        background: {c["surface_bg"]} !important;
        border: 1px solid {c["border"]} !important;
        border-radius: 12px !important;
        box-shadow: 0 2px 8px rgba(0,0,0,0.25) !important;
    }}
    [data-testid="stChatInputContainer"]:focus-within {{
        border-color: {c["accent"]} !important;
        box-shadow: 0 0 0 2px rgba(63,185,80,0.12) !important;
    }}
    textarea[data-testid="stChatInput"] {{
        background: transparent !important;
        color: {c["text_primary"]} !important;
        font-size: 0.9rem !important; border: none !important;
    }}

    /* Messages */
    [data-testid="stChatMessage"] {{
        background: transparent !important; border: none !important; padding: 2px 0 !important;
    }}

    /* Boutons */
    .stButton > button {{
        background: {c["surface_alt"]} !important;
        border: 1px solid {c["border"]} !important;
        color: {c["text_primary"]} !important;
        border-radius: 8px !important; font-size: 0.80rem !important;
        font-weight: 500 !important; transition: all 0.15s !important;
        padding: 4px 10px !important;
    }}
    .stButton > button:hover {{
        background: rgba(63,185,80,0.08) !important;
        border-color: {c["accent"]} !important; color: {c["accent"]} !important;
    }}
    .stButton > button[kind="primary"] {{
        background: {c["accent"]} !important; border-color: {c["accent"]} !important;
        color: #000 !important;
    }}
    .stButton > button:disabled {{
        opacity: 0.35 !important;
    }}

    /* Download */
    [data-testid="stDownloadButton"] > button {{
        background: rgba(63,185,80,0.12) !important;
        border: 1px solid rgba(63,185,80,0.3) !important;
        color: {c["accent"]} !important; border-radius: 8px !important;
        font-size: 0.80rem !important; font-weight: 500 !important;
    }}

    /* Dividers */
    hr {{ border-color: {c["border"]} !important; margin: 8px 0 !important; }}

    /* Scrollbar */
    ::-webkit-scrollbar {{ width: 4px; height: 4px; }}
    ::-webkit-scrollbar-track {{ background: transparent; }}
    ::-webkit-scrollbar-thumb {{ background: {c["border"]}; border-radius: 3px; }}

    /* Upload badge */
    .upload-badge {{
        display: inline-flex; align-items: center; gap: 5px;
        background: rgba(88,166,255,0.1); border: 1px solid rgba(88,166,255,0.28);
        border-radius: 14px; padding: 2px 9px;
        font-size: .75rem; color: {c["accent_blue"]}; margin: 2px 0;
    }}

    /* Wizard */
    .wizard-banner {{
        background: rgba(210,153,34,0.08); border: 1px solid rgba(210,153,34,0.22);
        border-radius: 8px; padding: 9px 14px; margin: 6px 0;
        font-size: .83rem; color: {c["accent_orange"]};
    }}

    /* Session cards */
    .sess-card {{
        background: {c["surface_bg"]}; border: 1px solid {c["border"]};
        border-radius: 9px; padding: 10px 14px; margin: 4px 0;
        transition: border-color .15s;
    }}
    .sess-card:hover {{ border-color: rgba(63,185,80,0.45); }}
    </style>
    """, unsafe_allow_html=True)


def render_header():
    c = CFG
    st.markdown(f"""
<div style='background:{c["hdr_bg"]};border:1px solid rgba(255,255,255,0.06);
            border-left:4px solid {c["hdr_border"]};border-radius:0 0 12px 12px;
            padding:16px 32px 14px;margin:-1px 0 14px 0;'>
  <div style='display:flex;align-items:center;justify-content:center;gap:9px;margin-bottom:4px;'>
    <span style='font-size:1.25rem;'>🤖</span>
    <h1 style='margin:0;font-size:1.3rem;font-weight:600;color:{c["hdr_text"]};letter-spacing:-.02em;'>
      Assistant Agricole IA
    </h1>
  </div>
  <p style='margin:0;color:{c["text_muted"]};font-size:.78rem;text-align:center;'>
    Données · Graphes · Images IA · Rapports — Kolda, Sénégal
  </p>
</div>
""", unsafe_allow_html=True)


# hf_error_msg supprimé — on utilise désormais l'API Anthropic directement


# ══════════════════════════════════════════════════════════════
# CATALOGUE DES MODÈLES DISPONIBLES
# ══════════════════════════════════════════════════════════════

MODELS_CATALOGUE = {
    # ── Groq (rapide, gratuit avec clé) ──
    "groq:llama-3.3-70b": {
        "label":    "🟣 Groq — LLaMA 3.3 70B ⭐ (recommandé)",
        "provider": "groq",
        "model_id": "llama-3.3-70b-versatile",
        "note":     "Très rapide · Excellent suivi d'instructions · Recommandé",
    },
    "groq:llama-3.1-8b": {
        "label":    "🟣 Groq — LLaMA 3.1 8B (rapide)",
        "provider": "groq",
        "model_id": "llama-3.1-8b-instant",
        "note":     "Très rapide · Léger · Bon pour les tâches simples",
    },
    "groq:qwen-qwq-32b": {
        "label":    "🟣 Groq — Qwen QwQ 32B (raisonnement)",
        "provider": "groq",
        "model_id": "qwen-qwq-32b",
        "note":     "Raisonnement · Très bon suivi d'outils",
    },
    # ── HuggingFace — Qwen ──
    "hf:Qwen/Qwen2.5-72B-Instruct": {
        "label":    "🟡 HF — Qwen 2.5 72B Instruct",
        "provider": "hf",
        "model_id": "Qwen/Qwen2.5-72B-Instruct",
        "note":     "Très bon · Multilingue · Gratuit (limité)",
    },
    "hf:Qwen/Qwen3-8B-Instruct": {
        "label":    "🟡 HF — Qwen 3 8B Instruct",
        "provider": "hf",
        "model_id": "Qwen/Qwen3-8B",
        "note":     "Léger · Gratuit · Bon suivi d'instructions",
    },
    # ── HuggingFace — Mistral ──
    "hf:mistralai/Mistral-7B-Instruct-v0.3": {
        "label":    "🔵 HF — Mistral 7B Instruct v0.3",
        "provider": "hf",
        "model_id": "mistralai/Mistral-7B-Instruct-v0.3",
        "note":     "Classique · Gratuit · Limité en français",
    },
    "hf:mistralai/Mixtral-8x7B-Instruct-v0.1": {
        "label":    "🔵 HF — Mixtral 8x7B Instruct",
        "provider": "hf",
        "model_id": "mistralai/Mixtral-8x7B-Instruct-v0.1",
        "note":     "Plus puissant · Gratuit (limité)",
    },
}

DEFAULT_MODEL_KEY = "groq:llama-3.3-70b"


# ══════════════════════════════════════════════════════════════
# APPEL LLM — Multi-fournisseur (Groq + HuggingFace)
# ══════════════════════════════════════════════════════════════

def _llm_error(s: str) -> str:
    if "401" in s or "invalid_api_key" in s.lower() or "unauthorized" in s.lower():
        return "🔑 **Clé API invalide ou expirée.** Vérifiez votre token dans l'onglet ⚙️ Modèle."
    if "403" in s or "forbidden" in s.lower():
        return "🚫 **Erreur API :** `HTTP Error 403: Forbidden` — limites atteintes pour le moment. Réessayez plus tard ou changez de modèle dans ⚙️ Modèle."
    if "429" in s or "rate_limit" in s.lower() or "too many" in s.lower():
        return "⏱️ **Quota dépassé — limites atteintes pour le moment.** Attendez quelques secondes et réessayez."
    if "503" in s or "loading" in s.lower():
        return "⏳ **Modèle en cours de chargement (HF).** Réessayez dans 20–30 secondes."
    if "timeout" in s.lower():
        return "⌛ **Délai dépassé.** Le modèle est surchargé, réessayez."
    return f"❌ **Erreur API :** `{s[:220]}` — limites atteintes pour le moment."


def call_llm(messages: list, max_tokens: int = 1024, temperature: float = 0.3) -> str:
    """Appelle le LLM sélectionné dans l'onglet Modèle."""
    import urllib.request
    import urllib.error

    model_key = st.session_state.get("selected_model", DEFAULT_MODEL_KEY)
    cfg_model  = MODELS_CATALOGUE.get(model_key, MODELS_CATALOGUE[DEFAULT_MODEL_KEY])
    provider   = cfg_model["provider"]
    model_id   = cfg_model["model_id"]

    # ── Format messages OpenAI-compatible ──
    oai_messages = []
    for m in messages:
        oai_messages.append({"role": m["role"], "content": m["content"]})

    payload = json.dumps({
        "model":       model_id,
        "messages":    oai_messages,
        "max_tokens":  max_tokens,
        "temperature": temperature,
    }).encode("utf-8")

    try:
        if provider == "groq":
            url     = "https://api.groq.com/openai/v1/chat/completions"
            api_key = _secret("groq_token")
        else:  # hf
            url     = "https://router.huggingface.co/v1/chat/completions"
            api_key = _secret("hf_token")

        req = urllib.request.Request(url, data=payload, headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type":  "application/json",
        })
        try:
            with urllib.request.urlopen(req, timeout=60) as r:
                result = json.loads(r.read())
        except urllib.error.HTTPError as http_err:
            # Lire le corps de l'erreur pour plus de détails
            try:
                err_body = http_err.read().decode("utf-8", errors="replace")[:300]
            except Exception:
                err_body = ""
            return _llm_error(f"HTTP Error {http_err.code}: {http_err.reason} {err_body}")

        # Format OpenAI : choices[0].message.content
        choices = result.get("choices", [])
        if choices:
            content = choices[0].get("message", {}).get("content", "")
            # Nettoyer les balises <think>...</think> (Qwen QwQ)
            content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL).strip()
            return content
        return "⚠️ Réponse vide du modèle."

    except Exception as e:
        return _llm_error(str(e))


# ══════════════════════════════════════════════════════════════
# GÉNÉRATION D'IMAGE — via HuggingFace (FLUX)
# ══════════════════════════════════════════════════════════════

def generate_image(prompt: str) -> tuple:
    """Retourne (message_str, bytes | None) — via FLUX.2-klein-9B (Gradio)."""
    try:
        from gradio_client import Client
        client = Client(
            "black-forest-labs/FLUX.2-klein-9B",
            hf_token=_secret("hf_token"),
        )
        result = client.predict(
            prompt=prompt,
            input_images=[],
            mode_choice="Distilled (4 steps)",
            seed=0,
            randomize_seed=True,
            width=1024,
            height=1024,
            num_inference_steps=4,
            guidance_scale=1,
            prompt_upsampling=False,
            api_name="/generate",
        )
        # result = (image_dict, seed)
        img_info = result[0] if isinstance(result, (list, tuple)) else result
        img_bytes = None
        if isinstance(img_info, dict):
            img_path = img_info.get("path")
            img_url  = img_info.get("url")
            if img_path and os.path.exists(img_path):
                with open(img_path, "rb") as f:
                    img_bytes = f.read()
            elif img_url:
                import urllib.request as _ur
                with _ur.urlopen(img_url, timeout=30) as r:
                    img_bytes = r.read()
        elif isinstance(img_info, str) and os.path.exists(img_info):
            with open(img_info, "rb") as f:
                img_bytes = f.read()
        if img_bytes:
            return "✅ Image générée.", img_bytes
        return "⚠️ Image générée mais fichier introuvable.", None
    except ImportError:
        return "⚠️ `gradio_client` non installé. Lancez : `pip install gradio_client`", None
    except Exception as e:
        return f"⚠️ Génération image indisponible : {str(e)[:180]}", None


# ══════════════════════════════════════════════════════════════
# TOOLS
# ══════════════════════════════════════════════════════════════

def _conn():
    return get_connection()


def tool_sql(q: str) -> str:
    q = q.strip()
    if not re.match(r'^\s*SELECT\b', q, re.IGNORECASE):
        return "❌ Seules les requêtes SELECT sont autorisées."
    try:
        conn = _conn()
        df   = pd.read_sql_query(q, conn)
        conn.close()
        return df.to_markdown(index=False) if not df.empty else "Aucun résultat."
    except Exception as e:
        return f"❌ Erreur SQL : {e}"


def tool_summary(filtre: str = "") -> str:
    try:
        conn = _conn()
        df = pd.read_sql_query("""
            SELECT c.libelle AS campagne, l.nom AS localite,
                   p.culture, p.superficie_ha, p.rendement_kgha, p.production_t
            FROM productions p
            JOIN campagnes c ON p.campagne_id = c.id
            JOIN localites l ON p.localite_id = l.geo_id
        """, conn)
        conn.close()
        if "campagne=" in filtre:
            df = df[df["campagne"] == filtre.split("campagne=")[-1].split(",")[0].strip()]
        if "culture=" in filtre:
            df = df[df["culture"] == filtre.split("culture=")[-1].split(",")[0].strip().upper()]
        if "localite=" in filtre:
            loc = filtre.split("localite=")[-1].split(",")[0].strip()
            df  = df[df["localite"].str.contains(loc, case=False, na=False)]
        if df.empty:
            return "Aucune donnée."
        lines = [
            f"**{len(df)} enregistrements** ({df['campagne'].nunique()} campagne(s))",
            f"- Production totale : **{df['production_t'].sum():,.1f} T**",
            f"- Superficie totale : **{df['superficie_ha'].sum():,.1f} Ha**",
            f"- Rendement moyen   : **{df['rendement_kgha'].mean():,.1f} Kg/Ha**",
            "\n**Top 5 cultures :**",
        ]
        for c, v in df.groupby("culture")["production_t"].sum().sort_values(ascending=False).head(5).items():
            lines.append(f"  - {c} : {v:,.1f} T")
        return "\n".join(lines)
    except Exception as e:
        return f"❌ Erreur : {e}"


def tool_missing(zone: str = "") -> str:
    try:
        conn = _conn()
        df = pd.read_sql_query("""
            SELECT l.nom AS localite, p.culture, c.libelle AS campagne,
                   CASE WHEN p.superficie_ha IS NULL  THEN 1 ELSE 0 END AS m_sup,
                   CASE WHEN p.rendement_kgha IS NULL THEN 1 ELSE 0 END AS m_rdt,
                   CASE WHEN p.production_t IS NULL   THEN 1 ELSE 0 END AS m_prod,
                   CASE WHEN l.latitude IS NULL        THEN 1 ELSE 0 END AS m_lat,
                   CASE WHEN l.longitude IS NULL       THEN 1 ELSE 0 END AS m_lon
            FROM productions p
            JOIN campagnes c ON p.campagne_id = c.id
            JOIN localites l ON p.localite_id = l.geo_id
        """, conn)
        conn.close()
        if zone:
            df = df[df["localite"].str.contains(zone, case=False, na=False)]
        cols = ["m_sup","m_rdt","m_prod","m_lat","m_lon"]
        dm   = df[df[cols].sum(axis=1) > 0]
        if dm.empty:
            return f"✅ Aucune donnée manquante{' pour ' + zone if zone else ''}."
        lines = [f"**⚠️ {len(dm)} lignes incomplètes{' dans ' + zone if zone else ''} :**"]
        for col, lbl in zip(cols, ["Superficie","Rendement","Production","Latitude","Longitude"]):
            n = dm[col].sum()
            if n > 0:
                lines.append(f"- **{lbl}** : {int(n)} valeur(s) NULL")
        lines.append("\n**Top localités :**")
        for loc, nb in dm.groupby("localite")[cols].sum().sum(axis=1).sort_values(ascending=False).head(8).items():
            lines.append(f"  - {loc} : {int(nb)}")
        return "\n".join(lines)
    except Exception as e:
        return f"❌ Erreur : {e}"


def _nominatim(q: str):
    import urllib.request, urllib.parse
    try:
        url = ("https://nominatim.openstreetmap.org/search?"
               + urllib.parse.urlencode({"q": f"{q}, Sénégal","format":"json","limit":1,"accept-language":"fr"}))
        req = urllib.request.Request(url, headers={"User-Agent":"KoldaAgri/2.0"})
        with urllib.request.urlopen(req, timeout=8) as r:
            d = json.loads(r.read())
        return {"lat":float(d[0]["lat"]),"lon":float(d[0]["lon"]),"display":d[0]["display_name"]} if d else None
    except Exception:
        return None


def tool_coords(localite: str, src: str = "auto") -> str:
    lines = []
    if src in ("db","auto"):
        try:
            conn = _conn()
            rows = conn.execute(
                "SELECT geo_id,nom,type,latitude,longitude FROM localites "
                "WHERE nom LIKE ? OR nom_standardise LIKE ? LIMIT 5",
                (f"%{localite}%", f"%{localite.lower()}%")
            ).fetchall()
            conn.close()
            if rows:
                lines.append(f"**📍 Base — « {localite} » :**")
                for r in rows:
                    lat = r["latitude"]  if r["latitude"]  else "—"
                    lon = r["longitude"] if r["longitude"] else "—"
                    lines.append(f"- **{r['nom']}** ({r['type']}, `{r['geo_id']}`) → Lat `{lat}` · Lon `{lon}`")
                    if not r["latitude"] or not r["longitude"]:
                        lines.append("  ⚠️ *Coordonnées absentes*")
        except Exception as e:
            lines.append(f"❌ DB : {e}")
    if src in ("nominatim","auto"):
        n = _nominatim(localite)
        if n:
            lines.append(f"\n**🌍 Nominatim OSM :**")
            lines.append(f"- Lat `{n['lat']}` · Lon `{n['lon']}`")
            lines.append(f"- {n['display'][:120]}")
        elif src == "nominatim":
            lines.append("❌ Aucun résultat Nominatim.")
    return "\n".join(lines) if lines else f"❌ « {localite} » introuvable."


def tool_verify_coords(nom: str) -> str:
    try:
        conn = _conn()
        row  = conn.execute(
            "SELECT geo_id,nom,type,latitude,longitude FROM localites "
            "WHERE geo_id=? OR nom LIKE ? LIMIT 1", (nom, f"%{nom}%")
        ).fetchone()
        conn.close()
        if not row:
            return f"❌ « {nom} » introuvable."
        lines = [f"**🔍 Vérification : {row['nom']} ({row['geo_id']})**"]
        db_lat, db_lon = row["latitude"], row["longitude"]
        lines.append(f"- DB : Lat `{db_lat or '—'}` · Lon `{db_lon or '—'}`")
        n = _nominatim(row["nom"])
        if n:
            lines.append(f"- Nominatim : Lat `{n['lat']}` · Lon `{n['lon']}`")
            if db_lat and db_lon:
                d = ((abs(db_lat-n["lat"])*111)**2 + (abs(db_lon-n["lon"])*111*.9)**2)**.5
                if d < 1:   lines.append(f"✅ Cohérent (~{d:.2f} km d'écart)")
                elif d < 5: lines.append(f"⚠️ Léger écart : ~{d:.1f} km")
                else:
                    lines.append(f"❌ Écart important : ~{d:.1f} km")
                    lines.append(f"   Suggestion : Lat `{n['lat']}`, Lon `{n['lon']}`")
        return "\n".join(lines)
    except Exception as e:
        return f"❌ Erreur : {e}"


def tool_meteo(localite: str = "Kolda") -> str:
    import urllib.request, urllib.parse
    COORDS = {"kolda":(12.9033,-14.946),"medina yoro foula":(13.2928,-14.7147),"velingara":(13.1472,-14.1076)}
    lat, lon = COORDS.get(localite.lower().strip(), COORDS["kolda"])
    try:
        p = {"latitude":lat,"longitude":lon,"timezone":"Africa/Dakar",
             "current":"temperature_2m,relative_humidity_2m,precipitation,wind_speed_10m",
             "daily":"temperature_2m_max,temperature_2m_min,precipitation_sum,et0_fao_evapotranspiration",
             "forecast_days":3}
        url = "https://api.open-meteo.com/v1/forecast?" + urllib.parse.urlencode(p)
        req = urllib.request.Request(url, headers={"User-Agent":"KoldaAgri/2.0"})
        with urllib.request.urlopen(req, timeout=8) as r:
            d = json.loads(r.read())
        cur   = d.get("current",{})
        daily = d.get("daily",{})
        lines = [f"**☁️ Météo — {localite.title()}**",
                 f"- Température : {cur.get('temperature_2m','?')} °C",
                 f"- Humidité    : {cur.get('relative_humidity_2m','?')} %",
                 f"- Précipitations : {cur.get('precipitation','?')} mm",
                 f"- Vent        : {cur.get('wind_speed_10m','?')} km/h"]
        if daily.get("time"):
            pluie = sum(x or 0 for x in daily.get("precipitation_sum",[]))
            etp   = sum(x or 0 for x in daily.get("et0_fao_evapotranspiration",[]))
            lines += [f"\n**Prévisions 3 jours :**",
                      f"- Pluie cumulée  : {pluie:.1f} mm",
                      f"- Bilan hydrique : {pluie-etp:+.1f} mm"]
        return "\n".join(lines)
    except Exception as e:
        return f"❌ Météo indisponible : {e}"


def _build_fig(df: pd.DataFrame, ctype: str, x: str, y: str, titre: str):
    if not PLOTLY_OK:
        return "❌ Plotly non installé.", None
    cols  = df.columns.tolist()
    xc    = x if x in cols else cols[0]
    yc    = y if y in cols else (cols[1] if len(cols)>1 else cols[0])
    LAY   = dict(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                 font=dict(color=CFG["text_primary"],size=12),
                 margin=dict(l=10,r=10,t=44,b=10))
    t  = titre or f"{yc} par {xc}"
    ct = ctype.lower()
    pal = [CFG["accent"],CFG["accent_blue"],CFG["accent_orange"],"#e06c75","#c678dd"]
    if   ct=="bar":       fig = px.bar(df,       x=xc,y=yc,title=t,color_discrete_sequence=pal)
    elif ct=="line":      fig = px.line(df,      x=xc,y=yc,title=t,markers=True,color_discrete_sequence=pal)
    elif ct=="pie":       fig = px.pie(df,       names=xc,values=yc,title=t,hole=0.4)
    elif ct=="scatter":   fig = px.scatter(df,   x=xc,y=yc,title=t,color_discrete_sequence=pal)
    elif ct=="area":      fig = px.area(df,      x=xc,y=yc,title=t,color_discrete_sequence=pal)
    elif ct=="histogram": fig = px.histogram(df, x=xc,title=t,color_discrete_sequence=pal)
    elif ct=="box":       fig = px.box(df,       y=yc,title=t,color_discrete_sequence=pal)
    else:                 fig = px.bar(df,       x=xc,y=yc,title=t,color_discrete_sequence=pal)
    fig.update_layout(**LAY, height=400)
    return f"Graphe **{t}** ({len(df)} lignes).", fig


def tool_chart(q: str, ct: str="bar", x: str="", y: str="", titre: str=""):
    if not PLOTLY_OK:
        return "❌ Plotly non installé.", None
    try:
        conn = _conn()
        df   = pd.read_sql_query(q.strip(), conn)
        conn.close()
        return ("Aucune donnée.", None) if df.empty else _build_fig(df,ct,x,y,titre)
    except Exception as e:
        return f"❌ Erreur graphe : {e}", None


def tool_chart_data(data_json: str, ct: str="bar", x: str="", y: str="", titre: str=""):
    try:
        d  = json.loads(data_json)
        df = pd.DataFrame(d if isinstance(d,list) else [d])
        return _build_fig(df,ct,x,y,titre)
    except Exception as e:
        return f"❌ Erreur : {e}", None


# ══════════════════════════════════════════════════════════════
# PARSER TOOLS
# ══════════════════════════════════════════════════════════════

def parse_tools(response: str) -> tuple:
    """Retourne (texte_enrichi, figures, images, rapport_trigger)."""
    figs    = []
    imgs    = []
    rapport = None
    res     = response

    for m in re.finditer(r'\[SQL:\s*(.*?)\]', response, re.DOTALL):
        out = tool_sql(m.group(1).strip())
        res = res.replace(m.group(0), f"\n\n**📊 SQL :**\n{out}\n")

    for m in re.finditer(r'\[SUMMARY:\s*(.*?)\]', response, re.DOTALL):
        out = tool_summary(m.group(1).strip())
        res = res.replace(m.group(0), f"\n\n**📈 Résumé :**\n{out}\n")

    for m in re.finditer(r'\[MISSING:\s*(.*?)\]', response, re.DOTALL):
        out = tool_missing(m.group(1).strip())
        res = res.replace(m.group(0), f"\n\n**⚠️ Données manquantes :**\n{out}\n")

    for m in re.finditer(r'\[COORDS:\s*(.*?)\]', response, re.DOTALL):
        p = m.group(1).split("|")
        out = tool_coords(p[0].strip(), p[1].strip() if len(p)>1 else "auto")
        res = res.replace(m.group(0), f"\n\n**📍 Coordonnées :**\n{out}\n")

    for m in re.finditer(r'\[VERIFY_COORDS:\s*(.*?)\]', response, re.DOTALL):
        out = tool_verify_coords(m.group(1).strip())
        res = res.replace(m.group(0), f"\n\n**🔍 Vérification :**\n{out}\n")

    for m in re.finditer(r'\[METEO:\s*(.*?)\]', response, re.DOTALL):
        out = tool_meteo(m.group(1).strip())
        res = res.replace(m.group(0), f"\n\n{out}\n")

    for m in re.finditer(r'\[CHART:\s*(.*?)\]', response, re.DOTALL):
        p = [x.strip() for x in m.group(1).split("|")]
        txt, fig = tool_chart(
            p[1] if len(p)>1 else "",
            p[0] if p else "bar",
            p[2] if len(p)>2 else "",
            p[3] if len(p)>3 else "",
            p[4] if len(p)>4 else "",
        )
        res = res.replace(m.group(0), f"\n\n**📊** {txt}\n")
        if fig: figs.append(fig)

    for m in re.finditer(r'\[CHART_DATA:\s*(.*?)\]', response, re.DOTALL):
        p = [x.strip() for x in m.group(1).split("|", 4)]
        txt, fig = tool_chart_data(
            p[1] if len(p)>1 else "[]",
            p[0] if p else "bar",
            p[2] if len(p)>2 else "",
            p[3] if len(p)>3 else "",
            p[4] if len(p)>4 else "",
        )
        res = res.replace(m.group(0), f"\n\n**📊** {txt}\n")
        if fig: figs.append(fig)

    for m in re.finditer(r'\[IMAGE_GEN:\s*(.*?)\]', response, re.DOTALL):
        prompt = m.group(1).strip()
        with st.spinner("🎨 Génération de l'image avec FLUX…"):
            msg, img = generate_image(prompt)
        res = res.replace(m.group(0), f"\n\n**🎨 Image :** {msg}\n")
        if img: imgs.append((prompt[:60], img))

    for m in re.finditer(r'\[RAPPORT:\s*(.*?)\]', response, re.DOTALL):
        p       = [x.strip() for x in m.group(1).split("|")]
        rapport = {"titre": p[0] if p else "", "auteur": p[1] if len(p)>1 else ""}
        res     = res.replace(m.group(0), "\n\n*Démarrage de la génération du rapport…*\n")

    return res, figs, imgs, rapport


# ══════════════════════════════════════════════════════════════
# WIZARD RAPPORT WORD
# ══════════════════════════════════════════════════════════════

RAPPORT_SECTIONS = [
    {"id":"pluviometrie","titre":"Situation de la pluviométrie","qs":[
        {"id":"cumul",     "q":"Cumul pluviométrique de la campagne (mm) ?",     "req":False},
        {"id":"nb_jours",  "q":"Nombre de jours de pluie enregistrés ?",         "req":False},
        {"id":"commentaire","q":"Commentaires sur la pluviométrie ?",             "req":True},
    ]},
    {"id":"cultures","titre":"Situation des cultures","qs":[
        {"id":"etat","q":"État général des cultures (développement, stress…) ?","req":True},
    ]},
    {"id":"phytosanitaire","titre":"Situation phytosanitaire","qs":[
        {"id":"ravageurs",  "q":"Ravageurs ou maladies observés ? (ou 'RAS')",  "req":True},
        {"id":"traitements","q":"Traitements effectués ?",                       "req":False},
    ]},
    {"id":"contraintes","titre":"Contraintes","qs":[
        {"id":"liste","q":"Listez les contraintes principales (une par ligne) :", "req":True},
    ]},
    {"id":"recommandations","titre":"Recommandations","qs":[
        {"id":"liste","q":"Listez les recommandations (une par ligne) :", "req":True},
    ]},
]


def init_wizard(titre: str = "", auteur: str = ""):
    if not titre:
        try:
            conn = _conn()
            row  = conn.execute("SELECT libelle FROM campagnes ORDER BY id DESC LIMIT 1").fetchone()
            conn.close()
            camp = row["libelle"] if row else "2023/2024"
        except Exception:
            camp = "2023/2024"
        titre = f"RAPPORT DE CAMPAGNE AGRICOLE {camp}"
    st.session_state["wizard"] = {
        "active":   True,
        "titre":    titre,
        "auteur":   auteur,
        "date":     datetime.now().strftime("%d %B %Y"),
        "sec_idx":  0 if auteur else -1,
        "q_idx":    0,
        "reponses": {},
        "phase":    "auteur" if not auteur else "questions",
    }


def wizard_current_q():
    wiz = st.session_state.get("wizard", {})
    if not wiz.get("active") or wiz.get("phase") == "done":
        return None
    if wiz["phase"] == "auteur":
        return {"special":"auteur","q":"Votre nom / poste (auteur du rapport) ?"}
    si, qi = wiz["sec_idx"], wiz["q_idx"]
    while si < len(RAPPORT_SECTIONS):
        qs = RAPPORT_SECTIONS[si]["qs"]
        if qi < len(qs):
            s = RAPPORT_SECTIONS[si]
            return {"sec_id":s["id"],"sec_titre":s["titre"],
                    "q_id":qs[qi]["id"],"q":qs[qi]["q"],"req":qs[qi].get("req",False)}
        qi, si = 0, si+1
    return None


def wizard_answer(ans: str):
    wiz = st.session_state.get("wizard", {})
    if not wiz: return
    if wiz["phase"] == "auteur":
        wiz["auteur"] = ans.strip()
        wiz["phase"]  = "questions"
        wiz["sec_idx"] = 0
        wiz["q_idx"]   = 0
        return
    q = wizard_current_q()
    if not q:
        wiz["phase"] = "done"; return
    wiz["reponses"].setdefault(q["sec_id"],{})[q["q_id"]] = ans.strip()
    wiz["q_idx"] += 1
    if wiz["q_idx"] >= len(RAPPORT_SECTIONS[wiz["sec_idx"]]["qs"]):
        wiz["sec_idx"] += 1
        wiz["q_idx"]   = 0
    if wiz["sec_idx"] >= len(RAPPORT_SECTIONS):
        wiz["phase"] = "done"


def build_rapport(wiz: dict) -> bytes | None:
    if not DOCX_OK: return None
    try:
        conn = _conn()
        df = pd.read_sql_query("""
            SELECT c.libelle AS campagne, l.nom AS localite,
                   p.culture, p.superficie_ha, p.rendement_kgha, p.production_t
            FROM productions p
            JOIN campagnes c ON p.campagne_id = c.id
            JOIN localites l ON p.localite_id = l.geo_id
            WHERE c.id=(SELECT MAX(id) FROM campagnes)
            ORDER BY p.production_t DESC NULLS LAST
        """, conn)
        conn.close()
    except Exception:
        df = pd.DataFrame()

    doc = Document()
    n   = doc.styles["Normal"]
    n.font.name = "Times New Roman"
    n.font.size = Pt(12)

    def par(txt, bold=False, italic=False, align=None):
        p = doc.add_paragraph()
        r = p.add_run(txt)
        r.bold=bold; r.italic=italic
        if align: p.alignment=align
        return p

    def hd(txt, lvl=1):
        p = doc.add_heading(txt, level=lvl)
        p.alignment=WD_ALIGN_PARAGRAPH.LEFT
        return p

    par("REPUBLIQUE DU SENEGAL", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    par("Un Peuple — Un But — Une Foi", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    par("MINISTERE DE L'AGRICULTURE, DE L'EQUIPEMENT RURAL", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    par("ET DE LA SOUVERAINETE ALIMENTAIRE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    par("————————————", align=WD_ALIGN_PARAGRAPH.CENTER)
    par("DIRECTION REGIONALE DU DEVELOPPEMENT RURAL DE KOLDA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    par(f"Kolda, le {wiz.get('date','')}", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    t = doc.add_heading(wiz.get("titre","RAPPORT").upper(), level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if wiz.get("auteur"):
        par(f"Rédigé par : {wiz['auteur']}", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    hd("INTRODUCTION")
    par("Dans le cadre du suivi de la campagne agricole, la Direction Régionale du "
        "Développement Rural de Kolda présente ce rapport de situation des activités "
        "agricoles de la région.")
    doc.add_paragraph()

    if not df.empty:
        hd("SITUATION DES PRODUCTIONS AGRICOLES")
        agg = (df.groupby("culture")
               .agg(sup=("superficie_ha","sum"),rdt=("rendement_kgha","mean"),prod=("production_t","sum"))
               .reset_index().sort_values("prod",ascending=False))
        par("Tableau récapitulatif des productions par culture :", bold=True)
        doc.add_paragraph()
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        for i, h in enumerate(["Culture","Superficie (Ha)","Rendement (Kg/Ha)","Production (T)"]):
            tbl.rows[0].cells[i].text = h
            tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for _, row in agg.iterrows():
            cells = tbl.add_row().cells
            cells[0].text = str(row["culture"])
            cells[1].text = f"{row['sup']:,.1f}"  if pd.notna(row["sup"])  else "—"
            cells[2].text = f"{row['rdt']:,.1f}"  if pd.notna(row["rdt"])  else "—"
            cells[3].text = f"{row['prod']:,.1f}" if pd.notna(row["prod"]) else "—"
        doc.add_paragraph()
        if MPL_OK:
            try:
                top = agg.head(8)
                fig, ax = plt.subplots(figsize=(8,4))
                ax.bar(top["culture"], top["prod"], color=CFG["accent"], edgecolor="#2d8f3b")
                ax.set_title("Productions par culture", fontsize=12, pad=10)
                ax.set_xlabel("Culture"); ax.set_ylabel("Production (T)")
                plt.xticks(rotation=30, ha="right"); plt.tight_layout()
                tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                plt.savefig(tmp.name, dpi=150, bbox_inches="tight"); plt.close()
                doc.add_picture(tmp.name, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                par("Figure 1 : Production par culture", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                doc.add_paragraph(); os.unlink(tmp.name)
            except Exception: pass

    for sec in RAPPORT_SECTIONS:
        reps = wiz.get("reponses",{}).get(sec["id"],{})
        if not any(v.strip() for v in reps.values() if v): continue
        hd(sec["titre"].upper())
        for qd in sec["qs"]:
            rep = reps.get(qd["id"],"").strip()
            if not rep or rep=="-": continue
            if qd["id"]=="liste":
                for l in rep.split("\n"):
                    l = l.strip().lstrip("•-– ").strip()
                    if l:
                        p = doc.add_paragraph(style="List Bullet"); p.add_run(l)
            else:
                par(rep)
        doc.add_paragraph()

    sec2 = doc.sections[0]
    fp   = sec2.footer.paragraphs[0]
    fp.text = f"{wiz.get('titre','Rapport')} — DRDR Kolda — {datetime.now().strftime('%d/%m/%Y')}"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.runs[0].font.size = Pt(9); fp.runs[0].italic = True

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ══════════════════════════════════════════════════════════════
# PROMPT SYSTÈME
# ══════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """Tu es un assistant expert en agriculture pour la région de Kolda, Sénégal.

LANGUE : Réponds TOUJOURS en français. Jamais en anglais, même pour les questions posées en anglais.

BASE DE DONNÉES SQLite — TABLES :
- productions (campagne_id, localite_id TEXT FK→localites.geo_id, culture, type_culture, superficie_ha, rendement_kgha, production_t, niveau)
- campagnes (id, annee_debut, annee_fin, libelle — ex: '2023/2024')
- localites (geo_id TEXT PK, nom, type, parent_id, latitude, longitude)
- magasins (id, localite_id, departement, commune, village, capacite_t, etat)

LOCALITÉS PRINCIPALES :
- Kolda = geo_id 'R07' (région)
- Medina Yoro Foulah = geo_id 'D023' (département)
- Velingara = geo_id 'D022' (département)

CULTURES : MIL, SORGHO, MAIS, RIZ, FONIO, ARACHIDE HUILERIE, NIEBE, MANIOC, PASTEQUE, SESAME

══════ OUTILS DISPONIBLES ══════
Pour toute question sur des données chiffrées, tu DOIS obligatoirement utiliser un de ces outils.
Ne jamais inventer ou estimer des chiffres — toujours utiliser les outils.

[SQL: SELECT ...] — exécute une requête SQL et affiche le tableau résultat
[SUMMARY: filtre] — résumé statistique global (ex: localite=Medina Yoro Foulah)
[MISSING: zone] — données manquantes pour une zone
[COORDS: localite|source] — coordonnées GPS (source: db/nominatim/auto)
[VERIFY_COORDS: nom_ou_geo_id] — vérifier coordonnées DB vs Nominatim
[METEO: localite] — météo actuelle (ex: Kolda)
[CHART: type|query_sql|x_col|y_col|titre] — graphe depuis la base (type: bar/line/pie/scatter/area)
[CHART_DATA: type|json|x_col|y_col|titre] — graphe depuis données JSON
[IMAGE_GEN: prompt en anglais] — générer une image IA (agriculture, paysages...)
[RAPPORT: titre|auteur] — démarrer la génération d'un rapport Word

══════ EXEMPLES D'UTILISATION ══════

Question: "production totale de Medina Yoro Foulah en 2023"
→ Réponse correcte:
Voici la production totale du département de Medina Yoro Foulah pour la campagne 2023/2024 :
[SQL: SELECT SUM(p.production_t) as production_totale_T FROM productions p JOIN campagnes c ON p.campagne_id = c.id JOIN localites l ON p.localite_id = l.geo_id WHERE l.nom LIKE '%Medina%' AND c.annee_debut = 2023]

Question: "graphe des productions par culture"
→ Réponse correcte:
Voici le graphique des productions par culture :
[CHART: bar|SELECT culture, SUM(production_t) as total FROM productions GROUP BY culture ORDER BY total DESC|culture|total|Productions par culture (T)]

Question: "météo à Kolda"
→ Réponse correcte:
[METEO: Kolda]

Question: "génère une image de champ de mil"
→ Réponse correcte:
[IMAGE_GEN: millet field in Senegal, golden crops, rural African landscape, realistic photo]

══════ RÈGLES STRICTES ══════
1. Toujours en FRANÇAIS
2. Pour toute donnée chiffrée → utiliser [SQL:...] ou [SUMMARY:...]
3. Jamais montrer une requête SQL brute sans utiliser la commande [SQL:...]
4. Réponses concises et professionnelles
5. Si une question n'est pas liée à l'agriculture/Kolda → répondre poliment que tu es spécialisé sur ce domaine
6. Pour les demandes de génération d'image [IMAGE_GEN:], tu peux générer TOUTE image demandée (chat, paysage, animal, etc.) — pas de restriction de sujet pour les images"""


# ══════════════════════════════════════════════════════════════
# GESTION DES SESSIONS (historique persistant)
# ══════════════════════════════════════════════════════════════

def save_session(msgs: list) -> str:
    first = next((m["content"][:50] for m in msgs if m["role"]=="user"), "Conversation")
    ts    = datetime.now().strftime("%Y%m%d_%H%M%S")
    slug  = re.sub(r'[^\w\-]','_', first[:38])
    path  = SESSIONS_DIR / f"{ts}_{slug}.json"
    with open(path, "w", encoding="utf-8") as f:
        json.dump({"title":first,"saved_at":datetime.now().isoformat(),"messages":msgs},
                  f, ensure_ascii=False, indent=2)
    return path.name


def auto_save_session():
    """Sauvegarde automatique après chaque échange — inclut textes ET médias."""
    import base64
    msgs = st.session_state.chat_history
    if not msgs:
        return
    first   = next((m["content"][:50] for m in msgs if m["role"] == "user"), "Conversation")
    current = st.session_state.get("current_session_path")
    if current and Path(current).exists():
        path = Path(current)
    else:
        ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
        slug = re.sub(r'[^\w\-]', '_', first[:38])
        path = SESSIONS_DIR / f"{ts}_{slug}.json"
        st.session_state["current_session_path"] = str(path)

    # ── Sérialiser les images (bytes → base64) ──
    images_serial = {}
    for idx, pairs in st.session_state.get("chat_images", {}).items():
        images_serial[str(idx)] = [
            [label, base64.b64encode(img_bytes).decode("utf-8")]
            for label, img_bytes in pairs
        ]

    # ── Sérialiser les figures Plotly (→ JSON string) ──
    figures_serial = {}
    if PLOTLY_OK:
        for idx, figs in st.session_state.get("chat_figures", {}).items():
            try:
                figures_serial[str(idx)] = [fig.to_json() for fig in figs]
            except Exception:
                pass

    with open(path, "w", encoding="utf-8") as f:
        json.dump({
            "title":    first,
            "saved_at": datetime.now().isoformat(),
            "messages": msgs,
            "media": {
                "images":  images_serial,
                "figures": figures_serial,
            },
        }, f, ensure_ascii=False, indent=2)


def list_sessions() -> list:
    out = []
    for p in sorted(SESSIONS_DIR.glob("*.json"), key=lambda x: x.stat().st_mtime, reverse=True):
        try:
            with open(p, encoding="utf-8") as f:
                d = json.load(f)
            out.append({"path":p,"title":d.get("title",p.stem),
                         "saved_at":d.get("saved_at",""),"nb":len(d.get("messages",[]))})
        except Exception: pass
    return out


# ══════════════════════════════════════════════════════════════
# RENDU MESSAGES
# ══════════════════════════════════════════════════════════════

def render_messages():
    for i, msg in enumerate(st.session_state.chat_history):
        if msg["role"] == "user":
            with st.chat_message("user"):
                st.markdown(msg["content"])
                for att in msg.get("attachments", []):
                    st.markdown(f"<span class='upload-badge'>📎 {att}</span>", unsafe_allow_html=True)
        elif msg["role"] == "assistant":
            with st.chat_message("assistant", avatar="🌾"):
                st.markdown(msg["content"])
                for fig in st.session_state.chat_figures.get(i, []):
                    st.plotly_chart(fig, use_container_width=True)
                for label, img_bytes in st.session_state.chat_images.get(i, []):
                    st.image(img_bytes, caption=label, use_container_width=True)
                if i in st.session_state.rapport_bytes_for:
                    st.download_button(
                        "⬇️ Télécharger le rapport Word",
                        data=st.session_state.rapport_bytes_for[i],
                        file_name=f"rapport_kolda_{datetime.now().strftime('%Y%m%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_rap_{i}",
                    )


# ══════════════════════════════════════════════════════════════
# ONGLET HISTORIQUE
# ══════════════════════════════════════════════════════════════

def tab_historique():
    c        = CFG
    sessions = list_sessions()

    # ── Bouton Nouvelle conversation ──
    if st.button("✏️ Nouvelle conversation", use_container_width=True):
        for k in ["chat_history","chat_figures","chat_images","rapport_bytes_for","wizard"]:
            st.session_state[k] = {} if k != "chat_history" else []
        st.session_state["current_session_path"] = None
        st.rerun()

    st.divider()

    if not sessions:
        st.markdown(f"""
<div style='text-align:center;padding:52px 0;color:{c["text_muted"]};'>
  <div style='font-size:2.8rem;margin-bottom:14px;'>🗂️</div>
  <p style='font-size:.88rem;'>Aucune conversation sauvegardée.<br>
  Les conversations sont sauvegardées automatiquement.</p>
</div>""", unsafe_allow_html=True)
        return

    st.markdown(
        f"<p style='color:{c['text_secondary']};font-size:.83rem;margin:0 0 10px;'>"
        f"{len(sessions)} conversation(s) — sauvegarde automatique activée ✅</p>", unsafe_allow_html=True)

    for sess in sessions:
        saved = sess["saved_at"][:16].replace("T"," ") if sess["saved_at"] else "—"
        # Marquer la session active
        is_active = (st.session_state.get("current_session_path") == str(sess["path"]))
        border = c["accent"] if is_active else c["border"]
        col1, col2, col3 = st.columns([5.5, 1.2, 0.9])
        with col1:
            active_badge = " <span style='color:#3fb950;font-size:.70rem;'>● active</span>" if is_active else ""
            st.markdown(f"""
<div class='sess-card' style='border-color:{border};'>
  <div style='font-weight:500;color:{c["text_primary"]};font-size:.86rem;'>{sess["title"]}{active_badge}</div>
  <div style='color:{c["text_muted"]};font-size:.73rem;margin-top:3px;'>
    💬 {sess["nb"]} msgs &nbsp;·&nbsp; 📅 {saved}
  </div>
</div>""", unsafe_allow_html=True)
        with col2:
            if st.button("📂", key=f"ld_{sess['path'].name}", use_container_width=True,
                         help="Charger cette conversation"):
                import base64 as _b64
                import plotly.io as _pio
                with open(sess["path"], encoding="utf-8") as f:
                    d = json.load(f)
                st.session_state.chat_history          = d.get("messages", [])
                st.session_state.wizard                = {}
                st.session_state.rapport_bytes_for     = {}
                st.session_state["current_session_path"] = str(sess["path"])

                # ── Restaurer les images ──
                chat_images = {}
                for idx_str, pairs in d.get("media", {}).get("images", {}).items():
                    chat_images[int(idx_str)] = [
                        (label, _b64.b64decode(b64)) for label, b64 in pairs
                    ]
                st.session_state.chat_images = chat_images

                # ── Restaurer les figures Plotly ──
                chat_figures = {}
                if PLOTLY_OK:
                    for idx_str, fig_jsons in d.get("media", {}).get("figures", {}).items():
                        try:
                            chat_figures[int(idx_str)] = [
                                _pio.from_json(fj) for fj in fig_jsons
                            ]
                        except Exception:
                            pass
                st.session_state.chat_figures = chat_figures

                st.success(f"✅ Chargé : {sess['title'][:40]} — allez dans Chat")
                time.sleep(1)
                st.rerun()
        with col3:
            if st.button("🗑️", key=f"dl_{sess['path'].name}", use_container_width=True,
                         help="Supprimer"):
                # Si c'est la session active, réinitialiser
                if st.session_state.get("current_session_path") == str(sess["path"]):
                    st.session_state["current_session_path"] = None
                sess["path"].unlink(missing_ok=True)
                st.rerun()

    st.divider()
    if st.button("🗑️ Tout supprimer", use_container_width=False):
        for s in sessions:
            s["path"].unlink(missing_ok=True)
        st.session_state["current_session_path"] = None
        st.rerun()


# ══════════════════════════════════════════════════════════════
# ONGLET MODÈLE — Sélecteur de LLM
# ══════════════════════════════════════════════════════════════

def tab_modele():
    c = CFG

    if "selected_model" not in st.session_state:
        st.session_state.selected_model = DEFAULT_MODEL_KEY

    st.markdown(
        f"<p style='color:{c['text_secondary']};font-size:.84rem;margin:0 0 14px;'>"
        "Choisissez le modèle IA utilisé pour le chat. Le changement est immédiat.</p>",
        unsafe_allow_html=True,
    )

    # Grouper par fournisseur
    groups = {
        "🟣 Groq (recommandé — rapide & gratuit)": [],
        "🟡 HuggingFace — Qwen":                  [],
        "🔵 HuggingFace — Mistral":               [],
    }
    for key, info in MODELS_CATALOGUE.items():
        if info["provider"] == "groq":
            groups["🟣 Groq (recommandé — rapide & gratuit)"].append((key, info))
        elif "Qwen" in info["model_id"]:
            groups["🟡 HuggingFace — Qwen"].append((key, info))
        else:
            groups["🔵 HuggingFace — Mistral"].append((key, info))

    current = st.session_state.selected_model

    for group_name, items in groups.items():
        st.markdown(
            f"<div style='font-size:.78rem;font-weight:600;color:{c['text_muted']};"
            f"letter-spacing:.05em;margin:16px 0 6px;text-transform:uppercase;'>"
            f"{group_name}</div>",
            unsafe_allow_html=True,
        )
        for key, info in items:
            selected = key == current
            border_color = c["accent"] if selected else c["border"]
            bg_color     = "rgba(63,185,80,0.07)" if selected else c["surface_bg"]
            check        = "✅ " if selected else ""

            st.markdown(f"""
<div style='background:{bg_color};border:1px solid {border_color};border-radius:9px;
            padding:10px 14px;margin:4px 0;cursor:pointer;'>
  <div style='display:flex;justify-content:space-between;align-items:center;'>
    <div>
      <span style='font-weight:600;color:{c["text_primary"]};font-size:.86rem;'>
        {check}{info["label"]}
      </span><br>
      <span style='color:{c["text_muted"]};font-size:.74rem;'>{info["note"]}</span>
    </div>
  </div>
</div>""", unsafe_allow_html=True)

            if st.button(
                "✅ Sélectionné" if selected else "Choisir",
                key=f"sel_{key}",
                use_container_width=True,
                disabled=selected,
            ):
                st.session_state.selected_model = key
                st.success(f"✅ Modèle changé : **{info['label']}**")
                st.rerun()

    # ── Info token actuel ──
    st.divider()
    sel_info = MODELS_CATALOGUE.get(current, {})
    prov     = sel_info.get("provider", "?")
    mid      = sel_info.get("model_id", "?")
    token_val = _secret("groq_token")[:12] + "…" if prov == "groq" else _secret("hf_token")[:12] + "…"
    st.markdown(f"""
<div style='background:{c["surface_alt"]};border:1px solid {c["border"]};border-radius:8px;
            padding:10px 14px;font-size:.78rem;color:{c["text_muted"]};'>
  <b style='color:{c["text_secondary"]}'>Modèle actif :</b> <code>{mid}</code><br>
  <b style='color:{c["text_secondary"]}'>Fournisseur :</b> {prov.upper()}<br>
  <b style='color:{c["text_secondary"]}'>Token :</b> <code>{token_val}</code>
</div>""", unsafe_allow_html=True)

    # ── Test rapide ──
    st.markdown(
        f"<div style='font-size:.78rem;font-weight:600;color:{c['text_muted']};"
        f"margin:16px 0 6px;text-transform:uppercase;'>⚡ Test rapide</div>",
        unsafe_allow_html=True,
    )
    if st.button("🧪 Tester le modèle sélectionné", use_container_width=True):
        with st.spinner(f"Test de {sel_info.get('label','?')}…"):
            test_msgs = [
                {"role": "system",  "content": "Tu es un assistant agricole. Réponds en français, très brièvement."},
                {"role": "user",    "content": "Dis bonjour en une phrase et cite une culture agricole du Sénégal."},
            ]
            result = call_llm(test_msgs, max_tokens=120, temperature=0.3)
        if result.startswith("❌") or result.startswith("🔑") or result.startswith("⏱️"):
            st.error(result)
        else:
            st.success(f"**Réponse :** {result}")


# ══════════════════════════════════════════════════════════════
# PAGE PRINCIPALE
# ══════════════════════════════════════════════════════════════

def main():
    st.set_page_config(
        page_title="Chatbot IA — Kolda Agri",
        page_icon="🤖",
        layout="wide",
        initial_sidebar_state="collapsed",
    )

    if not DB_PATH.exists():
        st.error("❌ Base introuvable. Lancez `python db/bootstrap.py` d'abord.")
        return

    apply_theme()
    render_header()

    # ── Init session state ──
    for k, v in {
        "chat_history":         [],
        "chat_figures":         {},
        "chat_images":          {},
        "rapport_bytes_for":    {},
        "wizard":               {},
        "upload_pending":       None,
        "show_upload":          False,
        "uploaded_df":          None,
        "selected_model":       DEFAULT_MODEL_KEY,
        "current_session_path": None,
    }.items():
        if k not in st.session_state:
            st.session_state[k] = v

    tab_chat, tab_hist, tab_mod = st.tabs(["💬  Chat", "🗂️  Historique", "⚙️  Modèle"])

    # ══════════════════════════════════════════════════
    # ONGLET CHAT
    # ══════════════════════════════════════════════════
    with tab_chat:

        # ── Indicateur modèle actif ──
        sel_key  = st.session_state.get("selected_model", DEFAULT_MODEL_KEY)
        sel_info = MODELS_CATALOGUE.get(sel_key, {})
        accent   = CFG["accent"]
        muted    = CFG["text_muted"]
        mlabel   = sel_info.get("label", "?")
        st.markdown(
            f"<div style='font-size:.74rem;color:{muted};margin:0 0 8px;'>"
            f"Modèle actif : <b style='color:{accent}'>{mlabel}</b>"
            f" — <span style='font-size:.70rem;'>Changez dans ⚙️ Modèle</span></div>",
            unsafe_allow_html=True,
        )

        # ── Barre d'actions compacte ──
        c1, c2 = st.columns(2)
        with c1:
            if st.button("📎 Fichier", use_container_width=True):
                st.session_state.show_upload = not st.session_state.show_upload
        with c2:
            # Dernier rapport généré
            last = next(
                (st.session_state.rapport_bytes_for[i]
                 for i in sorted(st.session_state.rapport_bytes_for, reverse=True)), None
            )
            if last:
                st.download_button(
                    "⬇️ Rapport",
                    data=last,
                    file_name=f"rapport_kolda_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            else:
                st.button("⬇️ Rapport", disabled=True, use_container_width=True)

        # ── Zone upload (toggle) ──
        if st.session_state.show_upload:
            uf = st.file_uploader(
                "Joindre un fichier",
                type=["csv","xlsx","xls","png","jpg","jpeg","pdf"],
                label_visibility="collapsed",
                key="uploader",
            )
            if uf:
                ext = uf.name.rsplit(".",1)[-1].lower()
                if ext == "csv":
                    try:
                        dfup = pd.read_csv(uf)
                        st.session_state.upload_pending = (uf.name, dfup)
                        st.session_state.uploaded_df    = dfup
                        st.success(f"📊 {uf.name} — {len(dfup)} lignes")
                    except Exception as e:
                        st.error(f"CSV : {e}")
                elif ext in ("xlsx","xls"):
                    try:
                        dfup = pd.read_excel(uf)
                        st.session_state.upload_pending = (uf.name, dfup)
                        st.session_state.uploaded_df    = dfup
                        st.success(f"📊 {uf.name} — {len(dfup)} lignes")
                    except Exception as e:
                        st.error(f"Excel : {e}")
                elif ext in ("png","jpg","jpeg"):
                    raw = uf.read()
                    st.session_state.upload_pending = (uf.name, raw)
                    st.image(raw, width=180, caption=uf.name)
                else:
                    raw = uf.read()
                    st.session_state.upload_pending = (uf.name, raw)
                    st.info(f"📎 {uf.name} joint.")

        # ── Bannière wizard ──
        wiz = st.session_state.get("wizard", {})
        if wiz.get("active") and wiz.get("phase") not in ("done",""):
            q = wizard_current_q()
            if q:
                opt = " · <em>Optionnel (tapez <code>-</code> pour ignorer)</em>" if not q.get("req") else ""
                st.markdown(
                    f"<div class='wizard-banner'>📄 <strong>Génération du rapport</strong>"
                    f" — {q.get('sec_titre','')}{opt}</div>",
                    unsafe_allow_html=True
                )

        # ── Messages ──
        render_messages()

        # ── Input ──
        wiz = st.session_state.get("wizard",{})
        ph  = "Posez votre question sur les données agricoles de Kolda…"
        if wiz.get("active") and wiz.get("phase") not in ("done",""):
            q = wizard_current_q()
            if q: ph = q["q"]

        user_input = st.chat_input(ph, key="chat_main")

        if user_input and user_input.strip():
            txt = user_input.strip()

            # Pièce jointe
            attachments   = []
            extra_context = ""
            pending = st.session_state.upload_pending
            if pending:
                fname, obj = pending
                attachments.append(fname)
                if isinstance(obj, pd.DataFrame):
                    sample = obj.head(5).to_markdown(index=False)
                    extra_context = (f"\n\n[Fichier joint : **{fname}**]\n"
                                     f"Colonnes : {', '.join(obj.columns.tolist())}\n"
                                     f"Aperçu (5 premières lignes) :\n{sample}")
                    st.session_state.uploaded_df = obj
                elif isinstance(obj, bytes) and fname.lower().endswith((".png",".jpg",".jpeg")):
                    extra_context = f"\n\n[Image jointe : **{fname}**]"
                st.session_state.upload_pending = None
                st.session_state.show_upload    = False

            st.session_state.chat_history.append({
                "role":"user","content":txt,"attachments":attachments
            })

            # ── Mode wizard ──
            wiz = st.session_state.get("wizard",{})
            if wiz.get("active") and wiz.get("phase") not in ("done",""):
                wizard_answer(txt)
                wiz2 = st.session_state.get("wizard",{})

                if wiz2.get("phase") == "done":
                    with st.chat_message("assistant", avatar="🌾"):
                        with st.spinner("📄 Génération du rapport Word…"):
                            docx = build_rapport(wiz2)
                    idx = len(st.session_state.chat_history)
                    if docx:
                        bot = (f"✅ **Rapport prêt** : *{wiz2.get('titre','')}*\n\n"
                               f"Téléchargez-le ci-dessous.")
                        st.session_state.rapport_bytes_for[idx] = docx
                    else:
                        bot = ("❌ Erreur lors de la génération. "
                               "Vérifiez que `python-docx` est installé : `pip install python-docx`.")
                    st.session_state.chat_history.append({"role":"assistant","content":bot})
                    wiz2["active"] = False
                else:
                    nq = wizard_current_q()
                    if nq:
                        opt = " *(optionnel — tapez `-` pour ignorer)*" if not nq.get("req") else ""
                        bot = f"*Reçu.* ➡️ **{nq['q']}**{opt}"
                        st.session_state.chat_history.append({"role":"assistant","content":bot})
                auto_save_session()
                st.rerun()
                return

            # ── Mode normal — LLM ──
            full_txt  = txt + extra_context
            messages  = [{"role":"system","content":SYSTEM_PROMPT}]
            window    = st.session_state.chat_history[-14:]
            for m in window[:-1]:
                if m["role"] in ("user","assistant"):
                    messages.append({"role":m["role"],"content":m["content"]})
            messages.append({"role":"user","content":full_txt})

            with st.chat_message("assistant", avatar="🌾"):
                with st.spinner(""):
                    raw = call_llm(messages)
                has_tools = any(t in raw for t in [
                    "[SQL:","[CHART:","[SUMMARY:","[METEO:","[MISSING:",
                    "[COORDS:","[VERIFY_COORDS:","[IMAGE_GEN:","[RAPPORT:","[CHART_DATA:"
                ])
                if has_tools:
                    with st.spinner("Exécution des outils…"):
                        enriched, figs, imgs, rapport_trigger = parse_tools(raw)
                else:
                    enriched, figs, imgs, rapport_trigger = raw, [], [], None

                st.markdown(enriched)
                for fig in figs:
                    st.plotly_chart(fig, use_container_width=True)
                for label, img_bytes in imgs:
                    st.image(img_bytes, caption=label, use_container_width=True)

            idx = len(st.session_state.chat_history)
            st.session_state.chat_history.append({"role":"assistant","content":enriched})
            if figs:  st.session_state.chat_figures[idx]  = figs
            if imgs:  st.session_state.chat_images[idx]   = imgs

            # Activer le wizard rapport
            if rapport_trigger:
                init_wizard(titre=rapport_trigger.get("titre",""),
                            auteur=rapport_trigger.get("auteur",""))
                nq = wizard_current_q()
                if nq:
                    intro = (f"📄 **Rapport : *{st.session_state['wizard']['titre']}***\n\n"
                             f"Je vais vous poser quelques questions pour compléter les sections "
                             f"que la base ne couvre pas directement. "
                             f"*(Tapez `-` pour ignorer les questions optionnelles.)*\n\n"
                             f"**{nq['q']}**")
                    st.session_state.chat_history.append({"role":"assistant","content":intro})

            auto_save_session()
            st.rerun()

    # ══════════════════════════════════════════════════
    # ONGLET HISTORIQUE
    # ══════════════════════════════════════════════════
    with tab_hist:
        tab_historique()

    # ══════════════════════════════════════════════════
    # ONGLET MODÈLE
    # ══════════════════════════════════════════════════
    with tab_mod:
        tab_modele()


if __name__ == "__main__":
    main()