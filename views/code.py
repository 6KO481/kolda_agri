"""
PAGE 6 — CHATBOT IA AVANCÉ
Agent conversationnel complet avec accès aux données agricoles de Kolda.

Fonctionnalités :
  1. Questions/réponses sur les données (SQL, données manquantes, coordonnées DB+Nominatim)
  2. Génération de graphes (depuis DB, données saisies ou fichiers uploadés)
     + génération d'images IA avec Flux (HuggingFace ZeroGPU)
  3. Mémoire conversationnelle + historique complet + modification de fichiers uploadés
  4. Génération de rapports Word (format rapport.docx) avec questions interactives
"""

import sys
import re
import io
import json
import time
import base64
import urllib.request
import urllib.parse
import urllib.error
import tempfile
import os
from pathlib import Path
from datetime import datetime

import streamlit as st
import pandas as pd

ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT / "db"))

from utils import get_connection, get_config, DB_PATH

try:
    import plotly.express as px
    import plotly.graph_objects as go
    PLOTLY_OK = True
except ImportError:
    PLOTLY_OK = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    MPL_OK = True
except ImportError:
    MPL_OK = False


# ══════════════════════════════════════════════════════════════
# THÈME & CONFIG
# ══════════════════════════════════════════════════════════════

@st.cache_data(ttl=60)
def load_config() -> dict:
    conn = get_connection()
    cfg  = get_config(conn)
    conn.close()
    return cfg

def cfg(key, default=""):
    return st.session_state.get("_config", {}).get(key, default)

def apply_theme():
    if "_config" not in st.session_state:
        st.session_state["_config"] = load_config()

    primary    = cfg("color_primary",      "#3fb950")
    font       = cfg("font_family",         "IBM Plex Mono, sans-serif").split(",")[0].strip()
    hdr_bg     = cfg("header_bg_color",     "#1c2a1e")
    hdr_border = cfg("header_border_color", "#3fb950")
    hdr_text   = cfg("header_text_color",   "#e6edf3")
    tab_active = cfg("tab_active_color",    "#3fb950")
    body_bg    = cfg("body_bg_color",       "#0d1117")

    def _hex_rgba(h, a=0.07):
        h = h.lstrip("#")
        if len(h) == 6:
            r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
            return f"rgba({r},{g},{b},{a})"
        return f"rgba(63,185,80,{a})"

    tab_bg = _hex_rgba(tab_active) if tab_active.startswith("#") else "rgba(63,185,80,0.07)"

    st.markdown(f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;500&family=Sora:wght@300;400;600;700&display=swap');
    html, body, [class*="css"] {{ font-family: '{font}', sans-serif !important; }}
    .main .block-container {{ padding-top: 0 !important; margin-top: 0 !important; }}
    header[data-testid="stHeader"] {{ height: 0 !important; }}
    .stApp {{ background-color: {body_bg} !important; }}
    .stButton > button[kind="primary"] {{
        background-color: {primary} !important; border-color: {primary} !important;
    }}
    [data-baseweb="tab-list"] {{ gap: 0 !important; width: 100% !important; }}
    [data-baseweb="tab"] {{
        flex: 1 1 0 !important; justify-content: center !important;
        padding: 10px 4px !important; font-size: 0.85rem !important;
        font-weight: 500 !important; border-bottom: 2px solid transparent !important;
    }}
    [data-baseweb="tab"][aria-selected="true"] {{
        border-bottom-color: {tab_active} !important;
        color: {tab_active} !important; background: {tab_bg} !important;
    }}
    .chat-tool {{
        background: rgba(88,166,255,0.08); border-left: 3px solid #58a6ff;
        border-radius: 0 6px 6px 0; padding: 6px 12px; margin: 4px 0;
        font-size: .8rem; color: #8b949e; font-family: monospace;
    }}
    .wizard-box {{
        background: rgba(210,153,34,0.1); border: 1px solid rgba(210,153,34,0.35);
        border-radius: 8px; padding: 14px 18px; margin: 10px 0;
    }}
    .info-pill {{
        display: inline-block; background: rgba(63,185,80,0.12);
        border: 1px solid rgba(63,185,80,0.3); border-radius: 20px;
        padding: 2px 10px; font-size: .78rem; color: #3fb950; margin: 2px;
    }}
    </style>
    """, unsafe_allow_html=True)
    return {"primary": primary, "hdr_bg": hdr_bg,
            "hdr_border": hdr_border, "hdr_text": hdr_text}


def render_header(theme):
    model = cfg("hf_model_id", "Mistral-7B").split("/")[-1]
    st.markdown(f"""
<div style='background:{theme["hdr_bg"]};border:1px solid rgba(255,255,255,0.06);
            border-left:4px solid {theme["hdr_border"]};
            border-radius:0 0 12px 12px;padding:18px 32px 16px;margin:-1px 0 20px 0;'>
  <div style='display:flex;align-items:center;justify-content:center;gap:10px;margin-bottom:6px;'>
    <span style='font-size:1.45rem;line-height:1;'>🤖</span>
    <h1 style='margin:0;font-size:1.45rem;font-weight:700;color:{theme["hdr_text"]};letter-spacing:-.01em;'>
      Assistant Agricole IA — Kolda
    </h1>
  </div>
  <p style='margin:0;color:#8b949e;font-size:.83rem;text-align:center;'>
    Données · Graphes · Images IA · Rapports Word &nbsp;|&nbsp; Modèle : <code>{model}</code>
  </p>
</div>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════
# TOOL 1 — SQL
# ══════════════════════════════════════════════════════════════

def tool_sql(query: str) -> str:
    query = query.strip()
    if not re.match(r'^\s*SELECT\b', query, re.IGNORECASE):
        return "❌ Seules les requêtes SELECT sont autorisées."
    try:
        conn = get_connection()
        df   = pd.read_sql_query(query, conn)
        conn.close()
        if df.empty:
            return "Aucun résultat pour cette requête."
        return df.to_markdown(index=False)
    except Exception as e:
        return f"❌ Erreur SQL : {e}"


# ══════════════════════════════════════════════════════════════
# TOOL 2 — DONNÉES MANQUANTES
# ══════════════════════════════════════════════════════════════

def tool_missing_data(zone: str = "") -> str:
    """Détecte les données manquantes (NULL) par zone/culture/campagne."""
    try:
        conn = get_connection()
        query = """
            SELECT c.libelle AS campagne, l.nom AS localite, l.type AS type_loc,
                   p.culture,
                   CASE WHEN p.superficie_ha   IS NULL THEN 1 ELSE 0 END AS manque_superficie,
                   CASE WHEN p.rendement_kgha   IS NULL THEN 1 ELSE 0 END AS manque_rendement,
                   CASE WHEN p.production_t     IS NULL THEN 1 ELSE 0 END AS manque_production,
                   CASE WHEN l.latitude         IS NULL THEN 1 ELSE 0 END AS manque_lat,
                   CASE WHEN l.longitude        IS NULL THEN 1 ELSE 0 END AS manque_lon
            FROM productions p
            JOIN campagnes c ON p.campagne_id = c.id
            JOIN localites l ON p.localite_id = l.geo_id
        """
        df = pd.read_sql_query(query, conn)
        conn.close()

        if zone.strip():
            df = df[df["localite"].str.contains(zone, case=False, na=False)]

        # Filtrer lignes avec au moins un manquant
        cols_manque = ["manque_superficie", "manque_rendement", "manque_production", "manque_lat", "manque_lon"]
        df_missing  = df[df[cols_manque].sum(axis=1) > 0].copy()

        if df_missing.empty:
            return f"✅ Aucune donnée manquante{' pour ' + zone if zone else ''} détectée."

        lines = [f"**⚠️ {len(df_missing)} lignes avec données manquantes{' dans ' + zone if zone else ''} :**\n"]

        # Résumé par type de manque
        for col, label in [
            ("manque_superficie", "Superficie (Ha)"),
            ("manque_rendement",  "Rendement (Kg/Ha)"),
            ("manque_production", "Production (T)"),
            ("manque_lat",        "Latitude localité"),
            ("manque_lon",        "Longitude localité"),
        ]:
            n = df_missing[col].sum()
            if n > 0:
                lines.append(f"- **{label}** : {int(n)} valeur(s) manquante(s)")

        # Détail par localité
        by_loc = (df_missing.groupby("localite")[cols_manque]
                  .sum().sum(axis=1).sort_values(ascending=False).head(10))
        lines.append("\n**Top localités concernées :**")
        for loc, nb in by_loc.items():
            lines.append(f"  - {loc} : {int(nb)} champ(s) incomplet(s)")

        return "\n".join(lines)
    except Exception as e:
        return f"❌ Erreur analyse données manquantes : {e}"


# ══════════════════════════════════════════════════════════════
# TOOL 3 — COORDONNÉES (DB + NOMINATIM)
# ══════════════════════════════════════════════════════════════

def _nominatim_search(query: str) -> dict | None:
    """Recherche Nominatim OSM. Retourne {'lat', 'lon', 'display_name'} ou None."""
    try:
        q = urllib.parse.quote_plus(f"{query}, Sénégal")
        url = f"https://nominatim.openstreetmap.org/search?q={q}&format=json&limit=1&accept-language=fr"
        req = urllib.request.Request(url, headers={"User-Agent": "KoldaAgri/2.0"})
        with urllib.request.urlopen(req, timeout=8) as r:
            data = json.loads(r.read())
        if data:
            return {"lat": float(data[0]["lat"]), "lon": float(data[0]["lon"]),
                    "display_name": data[0]["display_name"]}
        return None
    except Exception:
        return None


def tool_coords(localite: str, source: str = "auto") -> str:
    """
    Retourne les coordonnées d'une localité.
    source = 'db' | 'nominatim' | 'auto' (DB d'abord, puis Nominatim)
    """
    lines = []

    # --- Depuis la DB ---
    if source in ("db", "auto"):
        try:
            conn = get_connection()
            rows = conn.execute(
                "SELECT geo_id, nom, type, latitude, longitude FROM localites "
                "WHERE nom LIKE ? OR nom_standardise LIKE ? LIMIT 5",
                (f"%{localite}%", f"%{localite.lower()}%")
            ).fetchall()
            conn.close()
            if rows:
                lines.append(f"**📍 Résultats DB pour « {localite} » :**")
                for r in rows:
                    lat = r["latitude"]  if r["latitude"]  else "—"
                    lon = r["longitude"] if r["longitude"] else "—"
                    lines.append(f"- **{r['nom']}** ({r['type']}, {r['geo_id']}) → Lat: `{lat}`, Lon: `{lon}`")
                    if r["latitude"] is None or r["longitude"] is None:
                        lines.append(f"  ⚠️ *Coordonnées manquantes dans la base*")
        except Exception as e:
            lines.append(f"❌ Erreur DB : {e}")

    # --- Depuis Nominatim ---
    if source in ("nominatim", "auto"):
        nom_result = _nominatim_search(localite)
        if nom_result:
            lines.append(f"\n**🌍 Nominatim OSM :**")
            lines.append(f"- Latitude  : `{nom_result['lat']}`")
            lines.append(f"- Longitude : `{nom_result['lon']}`")
            lines.append(f"- Adresse   : {nom_result['display_name'][:120]}")
        elif source == "nominatim":
            lines.append("❌ Aucun résultat Nominatim pour cette localité.")

    return "\n".join(lines) if lines else f"❌ Localité « {localite} » introuvable."


def tool_verify_coords(geo_id_or_nom: str) -> str:
    """
    Compare les coordonnées de la DB avec celles de Nominatim.
    Signale les écarts importants (>0.05°, soit ~5km).
    """
    try:
        conn = get_connection()
        # Chercher par geo_id ou nom
        row = conn.execute(
            "SELECT geo_id, nom, type, latitude, longitude FROM localites "
            "WHERE geo_id = ? OR nom LIKE ? LIMIT 1",
            (geo_id_or_nom, f"%{geo_id_or_nom}%")
        ).fetchone()
        conn.close()

        if not row:
            return f"❌ Localité « {geo_id_or_nom} » introuvable dans la base."

        nom, geo_id = row["nom"], row["geo_id"]
        db_lat, db_lon = row["latitude"], row["longitude"]

        nom_result = _nominatim_search(nom)
        lines = [f"**🔍 Vérification coordonnées : {nom} ({geo_id})**"]

        if db_lat and db_lon:
            lines.append(f"- Base de données : Lat `{db_lat}`, Lon `{db_lon}`")
        else:
            lines.append("- Base de données : ⚠️ *Coordonnées absentes*")

        if nom_result:
            n_lat, n_lon = nom_result["lat"], nom_result["lon"]
            lines.append(f"- Nominatim OSM  : Lat `{n_lat}`, Lon `{n_lon}`")

            if db_lat and db_lon:
                delta_lat = abs(db_lat - n_lat)
                delta_lon = abs(db_lon - n_lon)
                dist_approx = ((delta_lat * 111) ** 2 + (delta_lon * 111 * 0.9) ** 2) ** 0.5
                if dist_approx < 1:
                    lines.append(f"✅ **Coordonnées cohérentes** (écart ~{dist_approx:.2f} km)")
                elif dist_approx < 5:
                    lines.append(f"⚠️ **Léger écart** : ~{dist_approx:.1f} km — À vérifier")
                else:
                    lines.append(f"❌ **Écart important** : ~{dist_approx:.1f} km — Coordonnées DB peut-être incorrectes")
                    lines.append(f"   → Suggestion : mettre à jour avec Lat `{n_lat}`, Lon `{n_lon}`")
        else:
            lines.append("- Nominatim : aucun résultat pour comparaison")

        return "\n".join(lines)
    except Exception as e:
        return f"❌ Erreur vérification : {e}"


# ══════════════════════════════════════════════════════════════
# TOOL 4 — RÉSUMÉ STATISTIQUE
# ══════════════════════════════════════════════════════════════

def tool_summary(filtre: str = "") -> str:
    try:
        conn = get_connection()
        df = pd.read_sql_query("""
            SELECT c.libelle AS campagne, l.nom AS localite,
                   p.culture, p.type_culture,
                   p.superficie_ha, p.rendement_kgha, p.production_t, p.niveau
            FROM productions p
            JOIN campagnes c ON p.campagne_id = c.id
            JOIN localites l ON p.localite_id = l.geo_id
        """, conn)
        conn.close()

        if "campagne=" in filtre:
            camp = filtre.split("campagne=")[-1].split(",")[0].strip()
            df = df[df["campagne"] == camp]
        if "culture=" in filtre:
            cult = filtre.split("culture=")[-1].split(",")[0].strip().upper()
            df = df[df["culture"] == cult]
        if "localite=" in filtre:
            loc = filtre.split("localite=")[-1].split(",")[0].strip()
            df = df[df["localite"].str.contains(loc, case=False, na=False)]
        if "niveau=" in filtre:
            niv = filtre.split("niveau=")[-1].split(",")[0].strip()
            df = df[df["niveau"] == niv]

        if df.empty:
            return "Aucune donnée pour ce filtre."

        lines = [f"**{len(df)} enregistrements** ({df['campagne'].nunique()} campagne(s))"]
        prod_total = df["production_t"].sum()
        sup_total  = df["superficie_ha"].sum()
        rdt_moy    = df["rendement_kgha"].mean()
        lines.append(f"- Production totale : **{prod_total:,.1f} T**")
        lines.append(f"- Superficie totale : **{sup_total:,.1f} Ha**")
        lines.append(f"- Rendement moyen   : **{rdt_moy:,.1f} Kg/Ha**")

        top = (df.groupby("culture")["production_t"].sum()
               .sort_values(ascending=False).head(5))
        lines.append("\n**Top 5 cultures :**")
        for cult, prod in top.items():
            lines.append(f"  - {cult} : {prod:,.1f} T")

        return "\n".join(lines)
    except Exception as e:
        return f"❌ Erreur : {e}"


# ══════════════════════════════════════════════════════════════
# TOOL 5 — MÉTÉO
# ══════════════════════════════════════════════════════════════

def tool_meteo(localite: str = "Kolda") -> str:
    COORDS = {
        "kolda":             (12.9033, -14.946),
        "medina yoro foula": (13.2928, -14.7147),
        "velingara":         (13.1472, -14.1076),
    }
    loc_key = localite.lower().strip()
    lat, lon = COORDS.get(loc_key, COORDS["kolda"])
    try:
        params = {
            "latitude":  lat, "longitude": lon,
            "timezone":  "Africa/Dakar",
            "current":   "temperature_2m,relative_humidity_2m,precipitation,wind_speed_10m",
            "daily":     "temperature_2m_max,temperature_2m_min,precipitation_sum,et0_fao_evapotranspiration",
            "forecast_days": 3,
        }
        url = "https://api.open-meteo.com/v1/forecast?" + urllib.parse.urlencode(params)
        req = urllib.request.Request(url, headers={"User-Agent": "KoldaAgri/2.0"})
        with urllib.request.urlopen(req, timeout=8) as r:
            data = json.loads(r.read())
        cur   = data.get("current", {})
        daily = data.get("daily", {})
        lines = [
            f"**Météo actuelle à {localite.title()}**",
            f"- Température : {cur.get('temperature_2m', '?')} °C",
            f"- Humidité    : {cur.get('relative_humidity_2m', '?')} %",
            f"- Précipitations : {cur.get('precipitation', '?')} mm",
            f"- Vent        : {cur.get('wind_speed_10m', '?')} km/h",
        ]
        if daily.get("time"):
            pluie_3j = sum(x or 0 for x in daily.get("precipitation_sum", []))
            etp_3j   = sum(x or 0 for x in daily.get("et0_fao_evapotranspiration", []))
            lines.append(f"\n**Prévisions 3 jours :**")
            lines.append(f"- Pluie cumulée : {pluie_3j:.1f} mm")
            lines.append(f"- ETP cumulée   : {etp_3j:.1f} mm")
            lines.append(f"- Bilan hydrique : {pluie_3j - etp_3j:+.1f} mm")
        return "\n".join(lines)
    except Exception as e:
        return f"❌ Météo indisponible ({e})."


# ══════════════════════════════════════════════════════════════
# TOOL 6 — GRAPHE DEPUIS LA DB
# ══════════════════════════════════════════════════════════════

def tool_chart(query: str, chart_type: str = "bar",
               x: str = "", y: str = "", titre: str = "") -> tuple:
    if not PLOTLY_OK:
        return "❌ Plotly non installé.", None
    try:
        conn = get_connection()
        df   = pd.read_sql_query(query.strip(), conn)
        conn.close()
        if df.empty:
            return "Aucune donnée pour ce graphe.", None
        return _build_plotly_fig(df, chart_type, x, y, titre)
    except Exception as e:
        return f"❌ Erreur graphe DB : {e}", None


def _build_plotly_fig(df: pd.DataFrame, chart_type: str,
                      x: str, y: str, titre: str) -> tuple:
    """Construit une figure Plotly depuis un DataFrame."""
    cols  = df.columns.tolist()
    x_col = x if x in cols else cols[0]
    y_col = y if y in cols else (cols[1] if len(cols) > 1 else cols[0])

    LAYOUT = dict(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#c9d1d9", size=12),
        margin=dict(l=10, r=10, t=45, b=10),
    )
    ct = chart_type.lower()
    t  = titre or f"{y_col} par {x_col}"
    if ct == "bar":
        fig = px.bar(df, x=x_col, y=y_col, title=t, color_discrete_sequence=["#3fb950"])
    elif ct == "line":
        fig = px.line(df, x=x_col, y=y_col, title=t, markers=True,
                      color_discrete_sequence=["#58a6ff"])
    elif ct == "pie":
        fig = px.pie(df, names=x_col, values=y_col, title=t, hole=0.4)
    elif ct == "scatter":
        fig = px.scatter(df, x=x_col, y=y_col, title=t,
                         color_discrete_sequence=["#d29922"])
    elif ct == "histogram":
        fig = px.histogram(df, x=x_col, title=t, color_discrete_sequence=["#3fb950"])
    elif ct == "box":
        fig = px.box(df, y=y_col, title=t, color_discrete_sequence=["#58a6ff"])
    elif ct == "area":
        fig = px.area(df, x=x_col, y=y_col, title=t,
                      color_discrete_sequence=["#3fb950"])
    else:
        fig = px.bar(df, x=x_col, y=y_col, title=t, color_discrete_sequence=["#3fb950"])

    fig.update_layout(**LAYOUT, height=400)
    return f"Graphe **{t}** généré ({len(df)} lignes).", fig


# ══════════════════════════════════════════════════════════════
# TOOL 7 — GRAPHE DEPUIS DONNÉES TEXTE / UPLOAD
# ══════════════════════════════════════════════════════════════

def tool_chart_from_data(data_json: str, chart_type: str = "bar",
                         x: str = "", y: str = "", titre: str = "") -> tuple:
    """
    Génère un graphe depuis des données JSON passées en texte.
    data_json : JSON d'un tableau ex: [{"culture":"MIL","prod":1200}, ...]
    """
    if not PLOTLY_OK:
        return "❌ Plotly non installé.", None
    try:
        data = json.loads(data_json)
        df   = pd.DataFrame(data) if isinstance(data, list) else pd.DataFrame([data])
        return _build_plotly_fig(df, chart_type, x, y, titre)
    except json.JSONDecodeError as e:
        return f"❌ JSON invalide : {e}", None
    except Exception as e:
        return f"❌ Erreur graphe données : {e}", None


def chart_from_uploaded_df(key: str = "uploaded_df",
                           chart_type: str = "bar",
                           x: str = "", y: str = "", titre: str = "") -> tuple:
    """Génère un graphe depuis le DataFrame uploadé en session_state."""
    df = st.session_state.get(key)
    if df is None:
        return "❌ Aucun fichier uploadé en session.", None
    return _build_plotly_fig(df, chart_type, x, y, titre)


# ══════════════════════════════════════════════════════════════
# TOOL 8 — GÉNÉRATION D'IMAGE IA (FLUX HF)
# ══════════════════════════════════════════════════════════════

def tool_image_gen(prompt: str, model: str = "black-forest-labs/FLUX.1-schnell") -> tuple:
    """
    Génère une image via HuggingFace Inference API (Flux).
    Retourne (message, bytes_image | None).
    """
    token = cfg("hf_api_token", "")
    if not token:
        return "❌ Token HuggingFace non configuré.", None

    # Enrichir le prompt avec contexte agricole sénégalais si pertinent
    payload = json.dumps({"inputs": prompt}).encode("utf-8")
    url     = f"https://api-inference.huggingface.co/models/{model}"
    req     = urllib.request.Request(
        url, data=payload,
        headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as r:
            img_bytes = r.read()
        # Vérifier que c'est bien une image
        if img_bytes[:4] in (b'\x89PNG', b'\xff\xd8\xff', b'RIFF', b'GIF8'):
            return f"✅ Image générée avec Flux (`{model.split('/')[-1]}`).", img_bytes
        # Peut-être du JSON d'erreur
        try:
            err = json.loads(img_bytes)
            if "error" in err:
                if "loading" in err["error"].lower():
                    return "⏳ Modèle en chargement, réessayez dans 30 secondes.", None
                return f"❌ Erreur API : {err['error']}", None
        except Exception:
            pass
        return "✅ Image générée.", img_bytes
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="ignore")
        if e.code == 503:
            return "⏳ Modèle Flux en chargement (ZeroGPU), réessayez dans 30s.", None
        return f"❌ Erreur HTTP {e.code} : {body[:150]}", None
    except Exception as e:
        return f"❌ Erreur génération image : {e}", None


# ══════════════════════════════════════════════════════════════
# TOOL 9 — GÉNÉRATION DE RAPPORT WORD
# ══════════════════════════════════════════════════════════════

# Sections standards du rapport agricole DRDR Kolda
RAPPORT_SECTIONS_DEF = [
    {
        "id": "pluviometrie",
        "titre": "Situation de la pluviométrie",
        "heading": 1,
        "questions": [
            {"id": "cumul_mm",       "q": "Quel est le cumul pluviométrique de la campagne (en mm) ?", "requis": False},
            {"id": "nb_jours_pluie", "q": "Combien de jours de pluie ont été enregistrés ?",           "requis": False},
            {"id": "commentaire",    "q": "Commentaires sur la pluviométrie (tendances, anomalies) ?",  "requis": True},
        ]
    },
    {
        "id": "cultures",
        "titre": "Situation des cultures",
        "heading": 1,
        "questions": [
            {"id": "commentaire",    "q": "Décrivez l'état général des cultures (développement végétatif, stress, etc.) :", "requis": True},
        ]
    },
    {
        "id": "phytosanitaire",
        "titre": "Situation phytosanitaire",
        "heading": 1,
        "questions": [
            {"id": "ravageurs",   "q": "Quels ravageurs ou maladies ont été observés ? (ou 'Aucun' si RAS)", "requis": True},
            {"id": "traitements", "q": "Quels traitements ont été effectués ?",                               "requis": False},
        ]
    },
    {
        "id": "contraintes",
        "titre": "Contraintes",
        "heading": 1,
        "questions": [
            {"id": "liste", "q": "Listez les principales contraintes de la campagne (une par ligne) :", "requis": True},
        ]
    },
    {
        "id": "recommandations",
        "titre": "Recommandations",
        "heading": 1,
        "questions": [
            {"id": "liste", "q": "Listez les recommandations principales (une par ligne) :", "requis": True},
        ]
    },
]


def _get_productions_for_rapport() -> pd.DataFrame:
    """Récupère les données de production de la dernière campagne."""
    try:
        conn = get_connection()
        df = pd.read_sql_query("""
            SELECT c.libelle AS campagne, l.nom AS localite, l.type AS type_loc,
                   p.culture, p.type_culture,
                   p.superficie_ha, p.rendement_kgha, p.production_t
            FROM productions p
            JOIN campagnes c ON p.campagne_id = c.id
            JOIN localites l ON p.localite_id = l.geo_id
            WHERE c.id = (SELECT MAX(id) FROM campagnes)
            ORDER BY p.production_t DESC NULLS LAST
        """, conn)
        conn.close()
        return df
    except Exception:
        return pd.DataFrame()


def generate_rapport_word(meta: dict, reponses: dict, df_prod: pd.DataFrame,
                          graphe_paths: list = None) -> bytes | None:
    """
    Génère un rapport Word au format DRDR Kolda.
    meta    : {titre, auteur, campagne, date}
    reponses: {section_id: {question_id: reponse}}
    df_prod : DataFrame des productions
    graphe_paths : liste de chemins d'images PNG à insérer
    Retourne les bytes du fichier .docx ou None si erreur.
    """
    if not DOCX_OK:
        return None

    doc = Document()

    # ── Style de base ──
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def add_para(text, bold=False, italic=False, centered=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        if centered:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def add_separator():
        p = doc.add_paragraph("─" * 60)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # ── En-tête institutionnel ──
    add_para("REPUBLIQUE DU SENEGAL", bold=True, centered=True)
    add_para("Un Peuple — Un But — Une Foi", italic=True, centered=True)
    add_para("", centered=True)
    add_para("MINISTERE DE L'AGRICULTURE, DE L'EQUIPEMENT RURAL", bold=True, centered=True)
    add_para("ET DE LA SOUVERAINETE ALIMENTAIRE", bold=True, centered=True)
    add_para("————————————————", centered=True)
    add_para("DIRECTION REGIONALE DU DEVELOPPEMENT RURAL DE KOLDA", bold=True, centered=True)
    add_para("", centered=True)
    add_para(f"Kolda, le {meta.get('date', datetime.now().strftime('%d %B %Y'))}", centered=True)
    doc.add_page_break()

    # ── Titre principal ──
    titre = meta.get("titre", f"RAPPORT DE CAMPAGNE AGRICOLE {meta.get('campagne', '')}")
    t_para = doc.add_heading(titre.upper(), level=0)
    t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if meta.get("auteur"):
        add_para(f"Rédigé par : {meta['auteur']}", italic=True, centered=True)
    add_para(f"Campagne : {meta.get('campagne', '—')}", centered=True)
    doc.add_page_break()

    # ── Introduction ──
    add_heading("INTRODUCTION", level=1)
    add_para(
        f"Dans le cadre du suivi de la campagne agricole {meta.get('campagne', '')}, "
        "la Direction Régionale du Développement Rural de Kolda présente le rapport "
        "de situation des activités agricoles de la région. Ce document fait le point "
        "sur les différentes composantes de la campagne : mise en place des facteurs "
        "de production, état des cultures, situation phytosanitaire, et les performances "
        "des différentes spéculations."
    )
    doc.add_paragraph()

    # ── Productions (depuis la base) ──
    if not df_prod.empty:
        add_heading("SITUATION DES PRODUCTIONS AGRICOLES", level=1)

        # Tableau récapitulatif par culture
        prod_tab = df_prod.groupby("culture").agg(
            superficie=("superficie_ha",   "sum"),
            rendement=("rendement_kgha",   "mean"),
            production=("production_t",    "sum"),
        ).reset_index().sort_values("production", ascending=False)

        add_para("Tableau récapitulatif des productions par culture :", bold=True)
        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(["Culture", "Superficie (Ha)", "Rendement (Kg/Ha)", "Production (T)"]):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        for _, row in prod_tab.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row["culture"])
            cells[1].text = f"{row['superficie']:,.1f}" if pd.notna(row["superficie"]) else "—"
            cells[2].text = f"{row['rendement']:,.1f}"  if pd.notna(row["rendement"])  else "—"
            cells[3].text = f"{row['production']:,.1f}" if pd.notna(row["production"]) else "—"

        doc.add_paragraph()

        # Totaux
        tot_sup  = df_prod["superficie_ha"].sum()
        tot_prod = df_prod["production_t"].sum()
        add_para(
            f"Au total, {tot_sup:,.1f} Ha ont été emblavés pour une production de "
            f"{tot_prod:,.1f} tonnes toutes cultures confondues.",
            bold=False
        )
        doc.add_paragraph()

    # ── Graphes (si disponibles) ──
    if graphe_paths:
        add_heading("GRAPHIQUES ET ANALYSES", level=1)
        for i, gpath in enumerate(graphe_paths):
            if gpath and os.path.exists(gpath):
                try:
                    doc.add_picture(gpath, width=Inches(5.5))
                    lp = doc.paragraphs[-1]
                    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_para(f"Figure {i+1}", italic=True, centered=True)
                    doc.add_paragraph()
                except Exception:
                    pass

    # ── Sections dynamiques (réponses du wizard) ──
    for sec_def in RAPPORT_SECTIONS_DEF:
        sec_id    = sec_def["id"]
        sec_reps  = reponses.get(sec_id, {})
        has_content = any(v.strip() for v in sec_reps.values() if v)
        if not has_content:
            continue

        add_heading(sec_def["titre"].upper(), level=sec_def["heading"])
        for q_def in sec_def["questions"]:
            qid = q_def["id"]
            rep = sec_reps.get(qid, "").strip()
            if not rep:
                continue
            if qid == "liste":
                for ligne in rep.split("\n"):
                    ligne = ligne.strip().lstrip("•-– ").strip()
                    if ligne:
                        p = doc.add_paragraph(style="List Bullet")
                        p.add_run(ligne)
            else:
                add_para(rep)
        doc.add_paragraph()

    # ── Pied de page ──
    section = doc.sections[0]
    footer  = section.footer
    fp = footer.paragraphs[0]
    fp.text = (f"{titre} — DRDR Kolda — "
               f"Généré le {datetime.now().strftime('%d/%m/%Y')} — Confidentiel")
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.runs[0].font.size = Pt(9)
    fp.runs[0].italic    = True

    # ── Sauvegarder en mémoire ──
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════
# WIZARD RAPPORT — GESTION ÉTAT
# ══════════════════════════════════════════════════════════════

def init_rapport_wizard(titre: str = "", auteur: str = "", campagne: str = ""):
    """Initialise le wizard de génération de rapport."""
    # Chercher la dernière campagne si non spécifiée
    if not campagne:
        try:
            conn = get_connection()
            row  = conn.execute("SELECT libelle FROM campagnes ORDER BY id DESC LIMIT 1").fetchone()
            conn.close()
            campagne = row["libelle"] if row else "2023/2024"
        except Exception:
            campagne = "2023/2024"

    st.session_state["rapport_wizard"] = {
        "active":           True,
        "titre":            titre or f"RAPPORT DE CAMPAGNE AGRICOLE {campagne}",
        "auteur":           auteur,
        "campagne":         campagne,
        "date":             datetime.now().strftime("%d %B %Y"),
        "sections_idx":     0,      # Index section courante dans RAPPORT_SECTIONS_DEF
        "question_idx":     0,      # Index question courante dans la section
        "reponses":         {},     # {section_id: {question_id: reponse}}
        "graphe_paths":     [],
        "phase":            "questions",   # 'questions' | 'done'
    }
    # Pré-remplir auteur si manquant
    if not st.session_state["rapport_wizard"]["auteur"]:
        st.session_state["rapport_wizard"]["phase"] = "ask_auteur"


def get_current_wizard_question() -> dict | None:
    """Retourne la question courante du wizard, ou None si terminé."""
    wiz = st.session_state.get("rapport_wizard")
    if not wiz or wiz["phase"] == "done":
        return None
    if wiz["phase"] == "ask_auteur":
        return {"special": "auteur", "q": "Quel est votre nom / poste (auteur du rapport) ?"}

    sec_idx  = wiz["sections_idx"]
    q_idx    = wiz["question_idx"]
    sections = RAPPORT_SECTIONS_DEF

    # Avancer jusqu'à une question non encore répondue
    while sec_idx < len(sections):
        sec  = sections[sec_idx]
        qs   = sec["questions"]
        while q_idx < len(qs):
            q_def = qs[q_idx]
            return {"sec_id": sec["id"], "sec_titre": sec["titre"],
                    "q_id": q_def["id"], "q": q_def["q"], "requis": q_def.get("requis", False),
                    "sec_idx": sec_idx, "q_idx": q_idx}
        q_idx   = 0
        sec_idx += 1

    return None  # Toutes les questions posées


def wizard_answer(answer: str):
    """Enregistre une réponse de wizard et avance à la prochaine question."""
    wiz = st.session_state.get("rapport_wizard")
    if not wiz:
        return

    if wiz["phase"] == "ask_auteur":
        wiz["auteur"] = answer.strip()
        wiz["phase"]  = "questions"
        return

    q_info = get_current_wizard_question()
    if not q_info:
        wiz["phase"] = "done"
        return

    sec_id = q_info["sec_id"]
    q_id   = q_info["q_id"]

    if sec_id not in wiz["reponses"]:
        wiz["reponses"][sec_id] = {}
    wiz["reponses"][sec_id][q_id] = answer.strip()

    # Avancer à la question suivante
    wiz["question_idx"] += 1
    sec   = RAPPORT_SECTIONS_DEF[wiz["sections_idx"]]
    if wiz["question_idx"] >= len(sec["questions"]):
        wiz["sections_idx"] += 1
        wiz["question_idx"]  = 0

    # Vérifier si on a terminé
    if wiz["sections_idx"] >= len(RAPPORT_SECTIONS_DEF):
        wiz["phase"] = "done"


def generate_rapport_and_respond() -> tuple:
    """Génère le rapport et retourne (message_chat, bytes_docx | None)."""
    wiz     = st.session_state.get("rapport_wizard")
    df_prod = _get_productions_for_rapport()

    # Générer graphe matplotlib des productions
    graphe_paths = []
    if MPL_OK and not df_prod.empty:
        try:
            top_cult = (df_prod.groupby("culture")["production_t"]
                        .sum().sort_values(ascending=False).head(8))
            fig, ax = plt.subplots(figsize=(9, 5))
            ax.bar(top_cult.index, top_cult.values, color="#3fb950", edgecolor="#2d8f3b")
            ax.set_title(f"Productions par culture — {wiz.get('campagne','')}", fontsize=13, pad=12)
            ax.set_xlabel("Culture")
            ax.set_ylabel("Production (T)")
            plt.xticks(rotation=35, ha="right")
            plt.tight_layout()
            tmpf = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            plt.savefig(tmpf.name, dpi=150, bbox_inches="tight")
            plt.close()
            graphe_paths.append(tmpf.name)
        except Exception:
            pass

    docx_bytes = generate_rapport_word(
        meta={
            "titre":    wiz.get("titre", ""),
            "auteur":   wiz.get("auteur", "DRDR Kolda"),
            "campagne": wiz.get("campagne", ""),
            "date":     wiz.get("date", ""),
        },
        reponses=wiz.get("reponses", {}),
        df_prod=df_prod,
        graphe_paths=graphe_paths,
    )

    # Nettoyer fichiers temporaires
    for p in graphe_paths:
        try:
            os.unlink(p)
        except Exception:
            pass

    # Désactiver le wizard
    st.session_state["rapport_wizard"]["active"] = False
    st.session_state["rapport_wizard"]["phase"]  = "done"

    if docx_bytes:
        msg = (f"✅ **Rapport généré** : *{wiz.get('titre', 'Rapport')}*\n\n"
               f"📄 Le fichier Word est prêt au téléchargement ci-dessous.\n\n"
               f"Le rapport comprend : en-tête institutionnel, tableau des productions, "
               f"graphique, et toutes les sections que vous avez renseignées.")
        return msg, docx_bytes
    else:
        return ("❌ Erreur lors de la génération du rapport Word. "
                "Vérifiez que `python-docx` est installé (`pip install python-docx`)."), None


# ══════════════════════════════════════════════════════════════
# AGENT LLM — APPEL HUGGINGFACE
# ══════════════════════════════════════════════════════════════

SYSTEM_PROMPT = """Tu es un assistant spécialisé dans l'agriculture de la région de Kolda au Sénégal.
Tu as accès à une base de données SQLite avec les tables :

TABLES :
- productions (campagne_id, localite_id, culture, type_culture, superficie_ha, rendement_kgha, production_t, niveau)
- campagnes (id, annee_debut, annee_fin, libelle)  — libelle ex: '2023/2024'
- localites (geo_id, nom, type, parent_id, latitude, longitude)
- magasins (id, localite_id, departement, commune, village, capacite_t, etat)

LOCALITÉS PRINCIPALES : Kolda (R07), Medina Yoro Foulah (D023), Velingara (D022)
CULTURES : MIL, SORGHO, MAIS, RIZ, FONIO, ARACHIDE HUILERIE, NIEBE, MANIOC, PASTEQUE, SESAME
NIVEAUX : 'localite', 'region'

COMMANDES DISPONIBLES (à inclure dans ta réponse) :

[SQL: SELECT ...] — requête SQL SELECT
[SUMMARY: filtre] — résumé statistique (filtre: campagne=X, culture=X, localite=X)
[MISSING: zone] — données manquantes (zone optionnelle, ex: Kolda)
[COORDS: localite|source] — coordonnées d'une localité (source: db/nominatim/auto)
[VERIFY_COORDS: nom_ou_geo_id] — vérifier si coordonnées DB sont exactes vs Nominatim
[METEO: localite] — météo actuelle
[CHART: type|query_sql|x_col|y_col|titre] — graphe depuis DB (type: bar/line/pie/scatter/area/box/histogram)
[CHART_DATA: type|json_data|x_col|y_col|titre] — graphe depuis données JSON
[IMAGE_GEN: prompt_en_anglais] — générer une image IA avec Flux
[RAPPORT: titre|auteur] — démarrer la génération d'un rapport Word interactif

RÈGLES :
- Réponds toujours en français
- Utilise les commandes pour interroger la base plutôt que d'inventer des données
- Pour les coordonnées manquantes, propose [COORDS: localite|nominatim]
- Pour les graphes, propose-les quand pertinent (CHART depuis DB, CHART_DATA depuis données saisies)
- Pour IMAGE_GEN: prompts en anglais, style réaliste ou agricole
- Si l'utilisateur demande un rapport, utilise [RAPPORT: titre|auteur]
- Signale proactivement les données manquantes ou incohérentes"""


def call_hf_api(messages: list, max_tokens: int = 1024, temperature: float = 0.3) -> str:
    token    = cfg("hf_api_token", "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")
    model_id = cfg("hf_model_id", "mistralai/Mistral-7B-Instruct-v0.3")

    if not token:
        return "❌ Token HuggingFace non configuré. Allez dans **Configuration → API & Chatbot**."

    # Format Mistral Instruct
    prompt = "<s>"
    for msg in messages:
        if msg["role"] == "system":
            prompt += f"[INST] {msg['content']} [/INST]\n"
        elif msg["role"] == "user":
            prompt += f"[INST] {msg['content']} [/INST]"
        elif msg["role"] == "assistant":
            prompt += f" {msg['content']}</s><s>"

    payload = {
        "inputs": prompt,
        "parameters": {
            "max_new_tokens":   max_tokens,
            "temperature":      max(temperature, 0.01),
            "return_full_text": False,
            "stop":             ["</s>", "[INST]"],
        }
    }
    url  = f"https://api-inference.huggingface.co/models/{model_id}"
    data = json.dumps(payload).encode("utf-8")
    req  = urllib.request.Request(
        url, data=data,
        headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    )
    try:
        with urllib.request.urlopen(req, timeout=40) as r:
            result = json.loads(r.read())
        if isinstance(result, list) and result:
            return result[0].get("generated_text", "").strip()
        elif isinstance(result, dict):
            if "error" in result:
                return f"❌ API Error: {result['error']}"
            return result.get("generated_text", "").strip()
        return "❌ Réponse API inattendue."
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="ignore")
        if e.code == 503:
            return "⏳ Le modèle est en cours de chargement, patientez 30s et réessayez."
        return f"❌ Erreur HTTP {e.code}: {body[:200]}"
    except Exception as e:
        return f"❌ Erreur réseau: {e}"


# ══════════════════════════════════════════════════════════════
# PARSER DES TOOLS
# ══════════════════════════════════════════════════════════════

def parse_and_execute_tools(response: str) -> tuple:
    """
    Parse la réponse LLM, exécute les tools, retourne (texte_enrichi, figures, images, rapport_trigger).
    """
    figures         = []
    images          = []       # [(label, bytes)]
    rapport_trigger = None
    result          = response

    # ── [SQL: ...] ──
    for match in re.finditer(r'\[SQL:\s*(.*?)\]', response, re.DOTALL):
        raw = match.group(1).strip()
        out = tool_sql(raw)
        result = result.replace(match.group(0), f"\n\n**📊 Résultat SQL :**\n{out}\n")

    # ── [SUMMARY: ...] ──
    for match in re.finditer(r'\[SUMMARY:\s*(.*?)\]', response, re.DOTALL):
        out    = tool_summary(match.group(1).strip())
        result = result.replace(match.group(0), f"\n\n**📈 Résumé statistique :**\n{out}\n")

    # ── [MISSING: ...] ──
    for match in re.finditer(r'\[MISSING:\s*(.*?)\]', response, re.DOTALL):
        out    = tool_missing_data(match.group(1).strip())
        result = result.replace(match.group(0), f"\n\n**⚠️ Données manquantes :**\n{out}\n")

    # ── [COORDS: localite|source] ──
    for match in re.finditer(r'\[COORDS:\s*(.*?)\]', response, re.DOTALL):
        parts = match.group(1).split("|")
        loc   = parts[0].strip()
        src   = parts[1].strip() if len(parts) > 1 else "auto"
        out   = tool_coords(loc, src)
        result = result.replace(match.group(0), f"\n\n**📍 Coordonnées :**\n{out}\n")

    # ── [VERIFY_COORDS: ...] ──
    for match in re.finditer(r'\[VERIFY_COORDS:\s*(.*?)\]', response, re.DOTALL):
        out    = tool_verify_coords(match.group(1).strip())
        result = result.replace(match.group(0), f"\n\n**🔍 Vérification coordonnées :**\n{out}\n")

    # ── [METEO: ...] ──
    for match in re.finditer(r'\[METEO:\s*(.*?)\]', response, re.DOTALL):
        out    = tool_meteo(match.group(1).strip())
        result = result.replace(match.group(0), f"\n\n**☁️ Météo :**\n{out}\n")

    # ── [CHART: type|query|x|y|titre] ──
    for match in re.finditer(r'\[CHART:\s*(.*?)\]', response, re.DOTALL):
        parts  = [p.strip() for p in match.group(1).split("|")]
        ctype  = parts[0] if len(parts) > 0 else "bar"
        cquery = parts[1] if len(parts) > 1 else ""
        cx     = parts[2] if len(parts) > 2 else ""
        cy     = parts[3] if len(parts) > 3 else ""
        ctitle = parts[4] if len(parts) > 4 else ""
        text_out, fig = tool_chart(cquery, ctype, cx, cy, ctitle)
        result = result.replace(match.group(0), f"\n\n**📊 Graphe :** {text_out}\n")
        if fig:
            figures.append(fig)

    # ── [CHART_DATA: type|json|x|y|titre] ──
    for match in re.finditer(r'\[CHART_DATA:\s*(.*?)\]', response, re.DOTALL):
        parts  = [p.strip() for p in match.group(1).split("|", 4)]
        ctype  = parts[0] if len(parts) > 0 else "bar"
        cjson  = parts[1] if len(parts) > 1 else "[]"
        cx     = parts[2] if len(parts) > 2 else ""
        cy     = parts[3] if len(parts) > 3 else ""
        ctitle = parts[4] if len(parts) > 4 else ""
        text_out, fig = tool_chart_from_data(cjson, ctype, cx, cy, ctitle)
        result = result.replace(match.group(0), f"\n\n**📊 Graphe :** {text_out}\n")
        if fig:
            figures.append(fig)

    # ── [IMAGE_GEN: prompt] ──
    for match in re.finditer(r'\[IMAGE_GEN:\s*(.*?)\]', response, re.DOTALL):
        prompt = match.group(1).strip()
        msg_out, img_bytes = tool_image_gen(prompt)
        result = result.replace(match.group(0), f"\n\n**🎨 Image IA :** {msg_out}\n")
        if img_bytes:
            images.append((prompt[:60], img_bytes))

    # ── [RAPPORT: titre|auteur] ──
    for match in re.finditer(r'\[RAPPORT:\s*(.*?)\]', response, re.DOTALL):
        parts  = [p.strip() for p in match.group(1).split("|")]
        titre  = parts[0] if len(parts) > 0 else ""
        auteur = parts[1] if len(parts) > 1 else ""
        rapport_trigger = {"titre": titre, "auteur": auteur}
        result = result.replace(match.group(0),
                                "\n\n*Je vais vous guider pour rédiger le rapport...*\n")

    return result, figures, images, rapport_trigger


# ══════════════════════════════════════════════════════════════
# INTERFACE PRINCIPALE
# ══════════════════════════════════════════════════════════════

QUESTIONS_RAPIDES = [
    ("🌾 Production totale",  "Quelle est la production totale de la région de Kolda pour la dernière campagne ?"),
    ("⚠️ Données manquantes", "Quelles sont les données manquantes dans la base pour la région de Kolda ?"),
    ("📍 Coordonnées Kolda",  "Quelles sont les coordonnées GPS de la région de Kolda dans la base ? Vérifie si elles sont correctes."),
    ("📊 Top cultures",       "Quelles sont les 5 cultures les plus produites à Kolda ? Génère un graphe."),
    ("💧 Météo actuelle",     "Quelle est la météo actuelle à Kolda et quel est son impact potentiel sur les cultures ?"),
    ("📈 Comparaison années", "Compare la production agricole entre les différentes campagnes disponibles dans la base."),
    ("🏪 Magasins",           "Quel est l'état des magasins de stockage dans la région de Kolda ?"),
    ("📄 Générer un rapport", "Je souhaite générer un rapport de campagne agricole pour Kolda."),
]


def main():
    st.set_page_config(
        page_title="Chatbot IA — Kolda Agri",
        page_icon="🤖",
        layout="wide",
    )

    if not DB_PATH.exists():
        st.error("❌ Base introuvable. Lancez `python db/bootstrap.py` d'abord.")
        return

    theme = apply_theme()
    render_header(theme)

    # ── Initialisation état session ─────────────────────────
    if "chat_history"  not in st.session_state:
        st.session_state.chat_history  = []
    if "chat_figures"  not in st.session_state:
        st.session_state.chat_figures  = {}
    if "chat_images"   not in st.session_state:
        st.session_state.chat_images   = {}
    if "rapport_bytes" not in st.session_state:
        st.session_state.rapport_bytes = None
    if "rapport_wizard" not in st.session_state:
        st.session_state.rapport_wizard = {}
    if "uploaded_df"   not in st.session_state:
        st.session_state.uploaded_df   = None

    # ── SIDEBAR ─────────────────────────────────────────────
    with st.sidebar:
        st.markdown("### ⚙️ Paramètres")
        max_tokens  = st.slider("Tokens max",  256, 2048, int(cfg("chatbot_max_tokens", "1024")), 64)
        temperature = st.slider("Température", 0.0, 1.0,  float(cfg("chatbot_temperature", "0.3")), 0.05)
        st.divider()

        # Token HF
        hf_token = cfg("hf_api_token", "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")
        if hf_token:
            st.success("✅ Token HuggingFace configuré")
        else:
            st.error("❌ Token manquant")
            st.page_link("views/5_Configuration.py", label="→ Configurer le token", icon="⚙️")

        st.divider()

        # ── Upload de fichier ──
        st.markdown("### 📂 Fichier de données")
        uploaded_file = st.file_uploader(
            "CSV / Excel pour graphe ou analyse",
            type=["csv", "xlsx", "xls"],
            help="Uploadez un fichier pour en générer des graphes ou l'analyser."
        )
        if uploaded_file:
            try:
                if uploaded_file.name.endswith(".csv"):
                    df_up = pd.read_csv(uploaded_file)
                else:
                    df_up = pd.read_excel(uploaded_file)
                st.session_state.uploaded_df = df_up
                st.success(f"✅ {len(df_up)} lignes · {len(df_up.columns)} colonnes")
                st.caption(f"Colonnes : {', '.join(df_up.columns[:6])}")
                if st.button("📊 Graphe auto depuis ce fichier", use_container_width=True):
                    cols = df_up.columns.tolist()
                    x = cols[0]
                    y = cols[1] if len(cols) > 1 else cols[0]
                    _, fig = _build_plotly_fig(df_up, "bar", x, y,
                                               f"Graphe — {uploaded_file.name}")
                    if fig:
                        idx = len(st.session_state.chat_history)
                        st.session_state.chat_history.append(
                            {"role": "assistant", "content": f"📊 Graphe généré depuis **{uploaded_file.name}**"})
                        st.session_state.chat_figures[idx] = [fig]
                        st.rerun()
            except Exception as e:
                st.error(f"Erreur lecture fichier : {e}")

        st.divider()

        # ── Rapport téléchargeable ──
        if st.session_state.rapport_bytes:
            st.markdown("### 📄 Rapport prêt")
            fname = f"rapport_kolda_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
            st.download_button(
                "⬇️ Télécharger le rapport Word",
                data=st.session_state.rapport_bytes,
                file_name=fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            if st.button("🗑️ Effacer le rapport", use_container_width=True):
                st.session_state.rapport_bytes = None
                st.rerun()
            st.divider()

        # ── Infos DB ──
        st.markdown("### 📊 Base de données")
        try:
            conn = get_connection()
            nb_prod = conn.execute("SELECT COUNT(*) FROM productions").fetchone()[0]
            nb_camp = conn.execute("SELECT COUNT(*) FROM campagnes").fetchone()[0]
            nb_loc  = conn.execute("SELECT COUNT(*) FROM localites").fetchone()[0]
            conn.close()
            st.caption(f"{nb_prod} productions · {nb_camp} campagnes · {nb_loc} localités")
        except Exception:
            pass

        st.divider()
        if st.button("🗑️ Vider la conversation", use_container_width=True):
            st.session_state.chat_history  = []
            st.session_state.chat_figures  = {}
            st.session_state.chat_images   = {}
            st.session_state.rapport_bytes = None
            st.session_state.rapport_wizard = {}
            st.rerun()

    # ── QUESTIONS RAPIDES (accueil) ──────────────────────────
    if not st.session_state.chat_history:
        st.markdown("#### 💬 Questions suggérées")
        cols = st.columns(4)
        for i, (label, question) in enumerate(QUESTIONS_RAPIDES):
            if cols[i % 4].button(label, key=f"qq_{i}", use_container_width=True):
                st.session_state.chat_history.append({"role": "user", "content": question})
                st.rerun()

    # ── AFFICHAGE HISTORIQUE ────────────────────────────────
    for i, msg in enumerate(st.session_state.chat_history):
        if msg["role"] == "user":
            with st.chat_message("user"):
                st.markdown(msg["content"])
        elif msg["role"] == "assistant":
            with st.chat_message("assistant", avatar="🌾"):
                st.markdown(msg["content"])
                # Figures Plotly
                for fig in st.session_state.chat_figures.get(i, []):
                    st.plotly_chart(fig, use_container_width=True)
                # Images IA
                for label, img_bytes in st.session_state.chat_images.get(i, []):
                    st.image(img_bytes, caption=label, use_container_width=True)

    # ── WIZARD RAPPORT (banner actif) ───────────────────────
    wiz = st.session_state.get("rapport_wizard", {})
    if wiz.get("active") and wiz.get("phase") not in ("done", ""):
        q_info = get_current_wizard_question()
        if q_info:
            sec_lbl = q_info.get("sec_titre", "Informations")
            with st.container():
                st.markdown(f"""
<div class='wizard-box'>
  <strong>📄 Génération de rapport — {wiz.get('titre','')[:60]}</strong><br>
  <span style='color:#8b949e;font-size:.85rem;'>Section : {sec_lbl}</span>
</div>
""", unsafe_allow_html=True)

    # ── ENTRÉE UTILISATEUR ──────────────────────────────────
    placeholder = "Posez votre question sur l'agriculture de Kolda…"
    wiz = st.session_state.get("rapport_wizard", {})
    if wiz.get("active") and wiz.get("phase") not in ("done", ""):
        q_info = get_current_wizard_question()
        if q_info:
            placeholder = q_info["q"]

    user_input = st.chat_input(placeholder, key="chat_input")

    if user_input and user_input.strip():
        txt = user_input.strip()
        st.session_state.chat_history.append({"role": "user", "content": txt})

        # ── Mode WIZARD ──
        wiz = st.session_state.get("rapport_wizard", {})
        if wiz.get("active") and wiz.get("phase") not in ("done", ""):
            q_info = get_current_wizard_question()

            # Enregistrer la réponse
            wizard_answer(txt)

            # Voir si on a terminé
            wiz_after = st.session_state.get("rapport_wizard", {})
            if wiz_after.get("phase") == "done":
                # Générer le rapport
                with st.chat_message("assistant", avatar="🌾"):
                    with st.spinner("📄 Génération du rapport Word en cours…"):
                        msg_out, docx_bytes = generate_rapport_and_respond()
                    st.markdown(msg_out)

                idx = len(st.session_state.chat_history)
                st.session_state.chat_history.append({"role": "assistant", "content": msg_out})
                if docx_bytes:
                    st.session_state.rapport_bytes = docx_bytes
            else:
                # Prochaine question
                next_q = get_current_wizard_question()
                if next_q:
                    bot_msg = (f"*Merci.* \n\n➡️ **{next_q['q']}**"
                               + (" *(optionnel — tapez `-` pour ignorer)*"
                                  if not next_q.get("requis") else ""))
                    with st.chat_message("assistant", avatar="🌾"):
                        st.markdown(bot_msg)
                    st.session_state.chat_history.append({"role": "assistant", "content": bot_msg})

            st.rerun()
            return

        # ── Mode NORMAL — appel LLM ──
        messages = [{"role": "system", "content": SYSTEM_PROMPT}]
        history_window = st.session_state.chat_history[-14:]
        messages.extend(history_window[:-1])
        messages.append({"role": "user", "content": txt})

        with st.chat_message("assistant", avatar="🌾"):
            with st.spinner("🤔 Réflexion en cours…"):
                raw_response = call_hf_api(messages, max_tokens=max_tokens,
                                           temperature=temperature)

            has_tools = any(tag in raw_response for tag in
                            ["[SQL:", "[CHART:", "[SUMMARY:", "[METEO:", "[MISSING:",
                             "[COORDS:", "[VERIFY_COORDS:", "[IMAGE_GEN:", "[RAPPORT:",
                             "[CHART_DATA:"])
            spin_ctx = st.spinner("🔧 Exécution des outils…") if has_tools else st.empty()
            with spin_ctx:
                enriched, figures, images, rapport_trigger = parse_and_execute_tools(raw_response)

            st.markdown(enriched)
            for fig in figures:
                st.plotly_chart(fig, use_container_width=True)
            for label, img_bytes in images:
                st.image(img_bytes, caption=label, use_container_width=True)

        # Sauvegarder
        idx = len(st.session_state.chat_history)
        st.session_state.chat_history.append({"role": "assistant", "content": enriched})
        if figures:
            st.session_state.chat_figures[idx] = figures
        if images:
            st.session_state.chat_images[idx] = images

        # Activer le wizard rapport si demandé
        if rapport_trigger:
            init_rapport_wizard(
                titre=rapport_trigger.get("titre", ""),
                auteur=rapport_trigger.get("auteur", ""),
            )
            next_q = get_current_wizard_question()
            if next_q:
                intro_msg = (
                    f"📄 **Démarrage de la génération du rapport : *{st.session_state['rapport_wizard']['titre']}***\n\n"
                    f"Je vais vous poser quelques questions pour compléter les informations "
                    f"que je ne trouve pas automatiquement dans la base. "
                    f"*(Pour ignorer une question optionnelle, tapez `-`)*\n\n"
                    f"**Première question :** {next_q['q']}"
                )
                with st.chat_message("assistant", avatar="🌾"):
                    st.markdown(intro_msg)
                st.session_state.chat_history.append({"role": "assistant", "content": intro_msg})

        st.rerun()

    # ── EXEMPLES (affichés vide) ─────────────────────────────
    if len(st.session_state.chat_history) == 0:
        with st.expander("💡 Exemples de questions que vous pouvez poser"):
            st.markdown("""
**📊 Données & analyses**
- *Quelle est la production de riz à Médina Yoro Foulah en 2023/2024 ?*
- *Montre-moi les données manquantes dans la zone de Kolda.*
- *Compare les rendements du mil entre 2020 et 2023 — génère un graphe.*

**📍 Géographie & coordonnées**
- *Quelle est la latitude et longitude de la région de Kolda dans la base ?*
- *Vérifie si les coordonnées de Vélingara dans la base sont exactes.*
- *Cherche les coordonnées de Médina Yoro Foulah via Nominatim.*

**🎨 Génération d'images IA**
- *Génère une image d'un champ de mil au Sénégal pendant l'hivernage.*
- *Crée une illustration d'un marché agricole en région de Kolda.*

**📄 Rapports Word**
- *Je veux générer un rapport de campagne agricole pour Kolda.*
- *Génère le rapport de la campagne 2023/2024 au format DRDR Kolda.*

**📂 Fichiers uploadés**
- *(Uploadez un CSV dans la sidebar)* → *Génère un graphe en barres avec ce fichier.*
""")


if __name__ == "__main__":
    main()